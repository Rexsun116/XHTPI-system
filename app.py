from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, get_flashed_messages, render_template_string, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
import os
from sqlalchemy import event, func, extract
from sqlalchemy.engine import Engine
from sqlalchemy.orm import joinedload
from sqlalchemy.orm.attributes import flag_modified
from docx import Document
from weasyprint import HTML
import requests

# 初始化app、db、login_manager
app = Flask(__name__)
app.secret_key = 'xhtpi-2024-secret-key'
# 正常运行仍使用 instance/database.db；测试/迁移演练可在启动前显式覆盖到隔离数据库。
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('XHTPI_DATABASE_URI', 'sqlite:///database.db')

# 仅供隔离测试/迁移演练显式启用；真实业务库目前仍保留现状，等待孤立 FK 人工处理。
if os.environ.get('XHTPI_ENABLE_SQLITE_FOREIGN_KEYS') == '1':
    @event.listens_for(Engine, 'connect')
    def enable_test_sqlite_foreign_keys(connection, _record):
        if connection.__class__.__module__.startswith('sqlite3'):
            cursor = connection.cursor()
            cursor.execute('PRAGMA foreign_keys=ON')
            cursor.close()

db = SQLAlchemy(app)
Migrate = Migrate(app, db)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# 自定义过滤器
@app.template_filter('format_number')
def format_number(value):
    """格式化数字，添加千位分隔符"""
    if value is None:
        return '0'
    try:
        return f"{int(value):,}"
    except (ValueError, TypeError):
        return str(value)

# 用户模型
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)
    role = db.Column(db.String(20), default='admin')
    
    def __init__(self, username=None, password_hash=None, role='admin'):
        self.username = username
        self.password_hash = password_hash
        self.role = role
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        if self.password_hash is None:
            return False
        return check_password_hash(self.password_hash, password)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@login_manager.unauthorized_handler
def unauthorized_callback():
    # 直接跳转到登录页，不闪现任何消息
    return redirect(url_for('login', next=request.path))

# 创建初始账号
@app.before_request
def create_admin_users():
    if not hasattr(app, '_admin_created'):
        users = [
            ('Rs10', 'Rs10!)xht'),
            ('Bw18', 'Bw18!*xht'),
            ('Dd16', 'Dd16!^xht'),
        ]
        for username, pwd in users:
            if not User.query.filter_by(username=username).first():
                user = User(username=username)
                user.set_password(pwd)
                db.session.add(user)
        db.session.commit()
        app._admin_created = True

# 登录视图，必须在所有@login_required视图之前
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('index'))
        else:
            flash('用户名或密码错误', 'danger')
    return render_template_string('''
    {% extends 'base.html' %}
    {% block content %}
    <div class="container mt-5" style="max-width:400px;">
      <div class="card shadow-sm">
        <div class="card-header bg-primary text-white">登录</div>
        <div class="card-body">
          <form method="post">
            <div class="mb-3">
              <label for="username" class="form-label">账号</label>
              <input type="text" class="form-control" id="username" name="username" required>
            </div>
            <div class="mb-3">
              <label for="password" class="form-label">密码</label>
              <input type="password" class="form-control" id="password" name="password" required>
            </div>
            <button type="submit" class="btn btn-primary w-100">登录</button>
          </form>
          {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
              <div class="mt-3">
                {% for category, message in messages %}
                  <div class="alert alert-{{ category }}">{{ message }}</div>
                {% endfor %}
              </div>
            {% endif %}
          {% endwith %}
        </div>
      </div>
    </div>
    {% endblock %}
    ''')

# -----------------------------
# 模型定义：客户 Customer
# -----------------------------

class Customer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(10), unique=True)
    name = db.Column(db.String(100), nullable=False)
    address = db.Column(db.String(200), nullable=False)
    tax_code = db.Column(db.String(100))
    country = db.Column(db.String(50), nullable=False)
    contact_person = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    email = db.Column(db.String(100))
    
    def __init__(self, code=None, name=None, address=None, tax_code=None, country=None, contact_person=None, phone=None, email=None):
        self.code = code
        self.name = name
        self.address = address
        self.tax_code = tax_code
        self.country = country
        self.contact_person = contact_person
        self.phone = phone
        self.email = email

#模型定义：出口商 Exporter

class Exporter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(10), unique=True)
    name = db.Column(db.String(100), nullable=False)
    address = db.Column(db.String(200), nullable=False)
    tax_code = db.Column(db.String(100))
    country = db.Column(db.String(50), nullable=False)
    contact_person = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    email = db.Column(db.String(100))
    
    def __init__(self, code=None, name=None, address=None, tax_code=None, country=None, contact_person=None, phone=None, email=None):
        self.code = code
        self.name = name
        self.address = address
        self.tax_code = tax_code
        self.country = country
        self.contact_person = contact_person
        self.phone = phone
        self.email = email

#模型定义：厂家 Factory

class Factory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(10), unique=True)
    name = db.Column(db.String(100), nullable=False)
    address = db.Column(db.String(200))
    tax_code = db.Column(db.String(100))
    country = db.Column(db.String(50), nullable=False)
    contact_person = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    email = db.Column(db.String(100))
    
    def __init__(self, code=None, name=None, address=None, tax_code=None, country=None, contact_person=None, phone=None, email=None):
        self.code = code
        self.name = name
        self.address = address
        self.tax_code = tax_code
        self.country = country
        self.contact_person = contact_person
        self.phone = phone
        self.email = email

#模型定义：货代 FreightForwarder

class FreightForwarder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(10), unique=True)
    name = db.Column(db.String(100), nullable=False)
    address = db.Column(db.String(200), nullable=False)
    tax_code = db.Column(db.String(100))
    country = db.Column(db.String(50), nullable=False)
    contact_person = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    email = db.Column(db.String(100))
    
    def __init__(self, code=None, name=None, address=None, tax_code=None, country=None, contact_person=None, phone=None, email=None):
        self.code = code
        self.name = name
        self.address = address
        self.tax_code = tax_code
        self.country = country
        self.contact_person = contact_person
        self.phone = phone
        self.email = email

#模型定义：货代报价 FreightQuote

class FreightQuote(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    freight_forwarder_id = db.Column(db.Integer, db.ForeignKey('freight_forwarder.id'), nullable=False)
    shipping_company = db.Column(db.String(100), nullable=False)  # 船公司
    departure_port = db.Column(db.String(100), nullable=False)    # 出发港
    destination_port = db.Column(db.String(100), nullable=False)  # 目的港
    route_type = db.Column(db.String(20), nullable=False)         # 直航或转运
    price = db.Column(db.Float, nullable=False)                   # 价格
    quote_date = db.Column(db.Date, nullable=False)               # 报价日期
    valid_until = db.Column(db.Date, nullable=False)              # 价格有效期至
    destination_14_days = db.Column(db.String(10))                # 目的港14天（有或无）
    remarks = db.Column(db.Text)                                  # 备注
    
    # 关联关系
    freight_forwarder = db.relationship('FreightForwarder', backref='quotes')
    
    def __init__(self, freight_forwarder_id=None, shipping_company=None, departure_port=None, 
                 destination_port=None, route_type=None, price=None, quote_date=None, 
                 valid_until=None, destination_14_days=None, remarks=None):
        self.freight_forwarder_id = freight_forwarder_id
        self.shipping_company = shipping_company
        self.departure_port = departure_port
        self.destination_port = destination_port
        self.route_type = route_type
        self.price = price
        self.quote_date = quote_date
        self.valid_until = valid_until
        self.destination_14_days = destination_14_days
        self.remarks = remarks

#模型定义：PI

class PI(db.Model):
    __tablename__ = 'pi'
    id = db.Column(db.Integer, primary_key=True)
    pi_no = db.Column(db.String(50), unique=True, nullable=False)
    pi_date = db.Column(db.Date, nullable=False)
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=False)
    # Commission / 工厂直发订单允许没有普通销售出口商；与现有 migration/SQLite schema 保持一致。
    exporter_id = db.Column(db.Integer, db.ForeignKey('exporter.id'), nullable=True)
    payment_terms = db.Column(db.String(100))
    loading_port = db.Column(db.String(100))
    destination_port = db.Column(db.String(100))
    bank = db.Column(db.Text) # 存放整段银行账户信息文本
    shipment_date = db.Column(db.Date)
    note = db.Column(db.String(200))
    status = db.Column(db.String(20), default='新建')
    # 历史 migration 06df15a0fb70 已存在的兼容字段；本阶段不解释或回填其业务值。
    pi_type = db.Column(db.String(50), nullable=True)

    # 快照字段 - 保存创建时的基础信息
    customer_name_snapshot = db.Column(db.String(100))  # 客户名称快照
    customer_address_snapshot = db.Column(db.String(200))  # 客户地址快照
    customer_tax_code_snapshot = db.Column(db.String(100))  # 客户税号快照
    customer_country_snapshot = db.Column(db.String(50))  # 客户国家快照
    customer_contact_snapshot = db.Column(db.String(100))  # 客户联系人快照
    customer_phone_snapshot = db.Column(db.String(50))  # 客户电话快照
    customer_email_snapshot = db.Column(db.String(100))  # 客户邮箱快照
    
    exporter_name_snapshot = db.Column(db.String(100))  # 出口商名称快照
    exporter_address_snapshot = db.Column(db.String(200))  # 出口商地址快照
    exporter_tax_code_snapshot = db.Column(db.String(100))  # 出口商税号快照
    exporter_country_snapshot = db.Column(db.String(50))  # 出口商国家快照
    exporter_contact_snapshot = db.Column(db.String(100))  # 出口商联系人快照
    exporter_phone_snapshot = db.Column(db.String(50))  # 出口商电话快照
    exporter_email_snapshot = db.Column(db.String(100))  # 出口商邮箱快照

    products = db.relationship('PIItem', backref='pi', cascade="all, delete-orphan")  # 👈 注意这里不用引号
    customer = db.relationship('Customer')
    exporter = db.relationship('Exporter', foreign_keys='PI.exporter_id')
    # 装运准备字段（可为空）
    freight_forwarder_id = db.Column(db.Integer, db.ForeignKey('freight_forwarder.id'))
    freight_forwarder = db.relationship('FreightForwarder')
    ocean_freight = db.Column(db.Float)
    container_date = db.Column(db.Date)
    # Phase 2B: naive local business datetime; legacy container_date remains compatible.
    container_loading_at = db.Column(db.DateTime, nullable=True)
    container_location = db.Column(db.String(100))
    driver_name = db.Column(db.String(100), nullable=True)
    driver_phone = db.Column(db.String(50), nullable=True)
    vehicle_number = db.Column(db.String(50), nullable=True)
    etd = db.Column(db.Date)
    eta = db.Column(db.Date)
    coo_required = db.Column(db.String(10))   # '需要' / '不需要'
    apta_required = db.Column(db.String(10))
    export_license_required = db.Column(db.String(10))
    customs_docs_required = db.Column(db.String(10))
    # Phase 2A 正式 Required Facts：三态 Boolean（True / False / 未确认 NULL）。
    coc_required = db.Column(db.Boolean, nullable=True)
    coa_required = db.Column(db.Boolean, nullable=True)
    original_bl_required = db.Column(db.Boolean, nullable=True)
    obd_electronic_required = db.Column(db.Boolean, nullable=True)
    insurance_original_required = db.Column(db.Boolean, nullable=True)
    insurance_electronic_required = db.Column(db.Boolean, nullable=True)
    original_documents_mail_required = db.Column(db.Boolean, nullable=True)
    telex_release_required = db.Column(db.Boolean, nullable=True)
    #7.2日新加
    other_documents = db.Column(db.String(200))
    # PI状态选择已发运时需填写的表单字段
    bill_of_lading = db.Column(db.String(100))
    shipping_company = db.Column(db.String(100))
    coa_status = db.Column(db.String(20))
    insurance_status = db.Column(db.String(20))
    document_shipping_status = db.Column(db.String(20))
    tracking_number = db.Column(db.String(100))
    batch_no = db.Column(db.String(200)) #多个批号用 ; 分隔
    actual_departure_date = db.Column(db.Date)  # 实际发运日期 7.4日新增


    # --- 托书相关字段 ---
    vessel_info = db.Column(db.String(100))
    booking_number = db.Column(db.String(100))
    container_type_quantity = db.Column(db.String(20))
    shipping_mark = db.Column(db.String(50))
    freight_term = db.Column(db.String(50))
    contract_number = db.Column(db.String(100))
    freight_clause = db.Column(db.String(200))
    waybill_option = db.Column(db.String(50))
    container_number = db.Column(db.String(50))
    seal_number = db.Column(db.String(50))
    container_type = db.Column(db.String(20))
    quantity_units = db.Column(db.Float)
    gross_weight = db.Column(db.Float)
    volume = db.Column(db.Float)
    vgm = db.Column(db.String(50))
    total_quantity = db.Column(db.Float)
    total_quantity_unit = db.Column(db.String(20))
    total_weight = db.Column(db.Float)
    total_weight_unit = db.Column(db.String(20))
    total_volume = db.Column(db.Float)
    total_volume_unit = db.Column(db.String(20))
    total_vgm = db.Column(db.String(50))
    
    #PI状态选择已到港时需要填写的字段
    actual_arrival_date = db.Column(db.Date)  # 实际到港日期
    payment_received = db.Column(db.String(10))  # 货款收齐：'已收齐' / '未收齐'
    # Phase 2C structured payment facts. Legacy payment_terms/payment_received remain compatible.
    currency = db.Column(db.String(10), nullable=True)
    advance_payment_percent = db.Column(db.Numeric(5, 2), nullable=True)
    advance_payment_amount = db.Column(db.Numeric(18, 2), nullable=True)
    advance_received_amount = db.Column(db.Numeric(18, 2), nullable=True)
    advance_received_at = db.Column(db.DateTime, nullable=True)
    balance_payment_amount = db.Column(db.Numeric(18, 2), nullable=True)
    balance_received_amount = db.Column(db.Numeric(18, 2), nullable=True)
    balance_received_at = db.Column(db.DateTime, nullable=True)
    freight_invoice_amount = db.Column(db.Float)  # 货代账单金额
    freight_invoice_confirmed = db.Column(db.String(10))  # 货代账单确认状态 '已确认' / '未确认'
    freight_invoice_issued = db.Column(db.String(10))  # 货代发票开具状态 '已开具' / '未开具'
    # Phase 2D dual-currency freight settlement facts. Legacy generic fields remain unchanged.
    freight_usd_bill_required = db.Column(db.Boolean, nullable=True)
    freight_usd_amount = db.Column(db.Numeric(18, 2), nullable=True)
    freight_usd_confirmed = db.Column(db.Boolean, nullable=True)
    freight_cny_bill_required = db.Column(db.Boolean, nullable=True)
    freight_cny_amount = db.Column(db.Numeric(18, 2), nullable=True)
    freight_cny_confirmed = db.Column(db.Boolean, nullable=True)
    freight_paid_at = db.Column(db.DateTime, nullable=True)
    telex_release = db.Column(db.String(10))  # PI电放状态 '已电放' / '未电放'
    settlement_documents_required = db.Column(db.String(10))  # 结汇文件 '需要' / '不需要'
    
    #PI状态选择已完成时需要填写的字段（新增）
    freight_payment_status = db.Column(db.String(10))  # 运费付款状态：'已付款' / '未付款'
    
    commission_factory_id = db.Column(db.Integer, db.ForeignKey('factory.id'))
    commission_exporter_id = db.Column(db.Integer, db.ForeignKey('exporter.id'))
    commission_amount = db.Column(db.Float)
    commission_rate = db.Column(db.Float)
    commission_status = db.Column(db.String(20))
    factory_sale_amount = db.Column(db.Float)
    
    def __init__(self, pi_no=None, pi_date=None, customer_id=None, exporter_id=None, payment_terms=None, 
                 loading_port=None, destination_port=None, bank=None, shipment_date=None, note=None, 
                 status='新建', customer_name_snapshot=None, customer_address_snapshot=None, 
                 customer_tax_code_snapshot=None, customer_country_snapshot=None, customer_contact_snapshot=None,
                 customer_phone_snapshot=None, customer_email_snapshot=None, exporter_name_snapshot=None,
                 exporter_address_snapshot=None, exporter_tax_code_snapshot=None, exporter_country_snapshot=None,
                 exporter_contact_snapshot=None, exporter_phone_snapshot=None, exporter_email_snapshot=None,
                 freight_forwarder_id=None, ocean_freight=None, container_date=None,
                 container_loading_at=None, container_location=None, driver_name=None,
                 driver_phone=None, vehicle_number=None,
                 etd=None, eta=None, coo_required=None, apta_required=None, export_license_required=None,
                 customs_docs_required=None, coc_required=None, coa_required=None, original_bl_required=None,
                 obd_electronic_required=None, insurance_original_required=None,
                 insurance_electronic_required=None, original_documents_mail_required=None,
                 telex_release_required=None, other_documents=None, bill_of_lading=None, shipping_company=None,
                 coa_status=None, insurance_status=None, document_shipping_status=None, tracking_number=None,
                 batch_no=None, actual_departure_date=None, vessel_info=None, booking_number=None,
                 container_type_quantity=None, shipping_mark=None, freight_term=None, contract_number=None,
                 freight_clause=None, waybill_option=None, container_number=None, seal_number=None,
                 container_type=None, quantity_units=None, gross_weight=None, volume=None, vgm=None,
                 total_quantity=None, total_quantity_unit=None, total_weight=None, total_weight_unit=None,
                 total_volume=None, total_volume_unit=None, total_vgm=None, actual_arrival_date=None,
                 payment_received=None, freight_invoice_amount=None, freight_invoice_confirmed=None,
                 freight_invoice_issued=None, telex_release=None, settlement_documents_required=None,
                 freight_usd_bill_required=None, freight_usd_amount=None, freight_usd_confirmed=None,
                 freight_cny_bill_required=None, freight_cny_amount=None, freight_cny_confirmed=None,
                 freight_paid_at=None,
                 currency=None, advance_payment_percent=None, advance_payment_amount=None,
                 advance_received_amount=None, advance_received_at=None,
                 balance_payment_amount=None, balance_received_amount=None, balance_received_at=None,
                 freight_payment_status=None,
                 commission_factory_id=None, commission_exporter_id=None, commission_amount=None, commission_rate=None, commission_status=None, factory_sale_amount=None):
        self.pi_no = pi_no
        self.pi_date = pi_date
        self.customer_id = customer_id
        self.exporter_id = exporter_id
        self.payment_terms = payment_terms
        self.loading_port = loading_port
        self.destination_port = destination_port
        self.bank = bank
        self.shipment_date = shipment_date
        self.note = note
        self.status = status
        self.customer_name_snapshot = customer_name_snapshot
        self.customer_address_snapshot = customer_address_snapshot
        self.customer_tax_code_snapshot = customer_tax_code_snapshot
        self.customer_country_snapshot = customer_country_snapshot
        self.customer_contact_snapshot = customer_contact_snapshot
        self.customer_phone_snapshot = customer_phone_snapshot
        self.customer_email_snapshot = customer_email_snapshot
        self.exporter_name_snapshot = exporter_name_snapshot
        self.exporter_address_snapshot = exporter_address_snapshot
        self.exporter_tax_code_snapshot = exporter_tax_code_snapshot
        self.exporter_country_snapshot = exporter_country_snapshot
        self.exporter_contact_snapshot = exporter_contact_snapshot
        self.exporter_phone_snapshot = exporter_phone_snapshot
        self.exporter_email_snapshot = exporter_email_snapshot
        self.freight_forwarder_id = freight_forwarder_id
        self.ocean_freight = ocean_freight
        self.container_date = container_date
        self.container_loading_at = container_loading_at
        self.container_location = container_location
        self.driver_name = driver_name
        self.driver_phone = driver_phone
        self.vehicle_number = vehicle_number
        self.etd = etd
        self.eta = eta
        self.coo_required = coo_required
        self.apta_required = apta_required
        self.export_license_required = export_license_required
        self.customs_docs_required = customs_docs_required
        self.coc_required = coc_required
        self.coa_required = coa_required
        self.original_bl_required = original_bl_required
        self.obd_electronic_required = obd_electronic_required
        self.insurance_original_required = insurance_original_required
        self.insurance_electronic_required = insurance_electronic_required
        self.original_documents_mail_required = original_documents_mail_required
        self.telex_release_required = telex_release_required
        self.other_documents = other_documents
        self.bill_of_lading = bill_of_lading
        self.shipping_company = shipping_company
        self.coa_status = coa_status
        self.insurance_status = insurance_status
        self.document_shipping_status = document_shipping_status
        self.tracking_number = tracking_number
        self.batch_no = batch_no
        self.actual_departure_date = actual_departure_date
        self.vessel_info = vessel_info
        self.booking_number = booking_number
        self.container_type_quantity = container_type_quantity
        self.shipping_mark = shipping_mark
        self.freight_term = freight_term
        self.contract_number = contract_number
        self.freight_clause = freight_clause
        self.waybill_option = waybill_option
        self.container_number = container_number
        self.seal_number = seal_number
        self.container_type = container_type
        self.quantity_units = quantity_units
        self.gross_weight = gross_weight
        self.volume = volume
        self.vgm = vgm
        self.total_quantity = total_quantity
        self.total_quantity_unit = total_quantity_unit
        self.total_weight = total_weight
        self.total_weight_unit = total_weight_unit
        self.total_volume = total_volume
        self.total_volume_unit = total_volume_unit
        self.total_vgm = total_vgm
        self.actual_arrival_date = actual_arrival_date
        self.payment_received = payment_received
        self.currency = currency
        self.advance_payment_percent = advance_payment_percent
        self.advance_payment_amount = advance_payment_amount
        self.advance_received_amount = advance_received_amount
        self.advance_received_at = advance_received_at
        self.balance_payment_amount = balance_payment_amount
        self.balance_received_amount = balance_received_amount
        self.balance_received_at = balance_received_at
        self.freight_invoice_amount = freight_invoice_amount
        self.freight_invoice_confirmed = freight_invoice_confirmed
        self.freight_invoice_issued = freight_invoice_issued
        self.freight_usd_bill_required = freight_usd_bill_required
        self.freight_usd_amount = freight_usd_amount
        self.freight_usd_confirmed = freight_usd_confirmed
        self.freight_cny_bill_required = freight_cny_bill_required
        self.freight_cny_amount = freight_cny_amount
        self.freight_cny_confirmed = freight_cny_confirmed
        self.freight_paid_at = freight_paid_at
        self.telex_release = telex_release
        self.settlement_documents_required = settlement_documents_required
        self.freight_payment_status = freight_payment_status
        # 新增工厂直发订单相关字段
        self.commission_factory_id = commission_factory_id
        self.commission_exporter_id = commission_exporter_id
        self.commission_amount = commission_amount
        self.commission_rate = commission_rate
        self.commission_status = commission_status
        self.factory_sale_amount = factory_sale_amount

    commission_factory = db.relationship('Factory', foreign_keys='PI.commission_factory_id')
    commission_exporter = db.relationship('Exporter', foreign_keys='PI.commission_exporter_id')


class PIItem(db.Model):
    __tablename__ = 'pi_item'
    id = db.Column(db.Integer, primary_key=True)
    pi_id = db.Column(db.Integer, db.ForeignKey('pi.id'), nullable=False)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    factory_id = db.Column(db.Integer, db.ForeignKey('factory.id'), nullable=False)
    trade_term = db.Column(db.String(50))
    unit_price = db.Column(db.Float)
    quantity = db.Column(db.Float)
    total_price = db.Column(db.Float)
    
    # 快照字段 - 保存创建时的产品和厂家信息
    product_category_snapshot = db.Column(db.String(100))  # 产品类别快照
    product_brand_snapshot = db.Column(db.String(100))  # 产品品牌快照
    product_model_snapshot = db.Column(db.String(100))  # 产品型号快照
    product_packaging_snapshot = db.Column(db.String(100))  # 产品包装快照
    
    factory_name_snapshot = db.Column(db.String(100))  # 厂家名称快照
    factory_address_snapshot = db.Column(db.String(200))  # 厂家地址快照
    factory_tax_code_snapshot = db.Column(db.String(100))  # 厂家税号快照
    factory_country_snapshot = db.Column(db.String(50))  # 厂家国家快照
    factory_contact_snapshot = db.Column(db.String(100))  # 厂家联系人快照
    factory_phone_snapshot = db.Column(db.String(50))  # 厂家电话快照
    factory_email_snapshot = db.Column(db.String(100))  # 厂家邮箱快照
    
    product = db.relationship('Product')
    factory = db.relationship('Factory')
    
    def __init__(self, pi_id=None, product_id=None, factory_id=None, trade_term=None, unit_price=None, 
                 quantity=None, total_price=None, product_category_snapshot=None, product_brand_snapshot=None,
                 product_model_snapshot=None, product_packaging_snapshot=None, factory_name_snapshot=None,
                 factory_address_snapshot=None, factory_tax_code_snapshot=None, factory_country_snapshot=None,
                 factory_contact_snapshot=None, factory_phone_snapshot=None, factory_email_snapshot=None):
        self.pi_id = pi_id
        self.product_id = product_id
        self.factory_id = factory_id
        self.trade_term = trade_term
        self.unit_price = unit_price
        self.quantity = quantity
        self.total_price = total_price
        self.product_category_snapshot = product_category_snapshot
        self.product_brand_snapshot = product_brand_snapshot
        self.product_model_snapshot = product_model_snapshot
        self.product_packaging_snapshot = product_packaging_snapshot
        self.factory_name_snapshot = factory_name_snapshot
        self.factory_address_snapshot = factory_address_snapshot
        self.factory_tax_code_snapshot = factory_tax_code_snapshot
        self.factory_country_snapshot = factory_country_snapshot
        self.factory_contact_snapshot = factory_contact_snapshot
        self.factory_phone_snapshot = factory_phone_snapshot
        self.factory_email_snapshot = factory_email_snapshot

# 模型定义：产品 Product

class Product(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(10), unique=True)
    category = db.Column(db.String(100), nullable=False, default="Titanium dioxide")
    brand = db.Column(db.String(100))  # 可空，如 "TITAX" 或空
    model = db.Column(db.String(100), nullable=False)  # 型号，如 R-504
    packaging = db.Column(db.String(100))  # 包装规格，如 25KG/包
    
    def __init__(self, code=None, category="Titanium dioxide", brand=None, model=None, packaging=None):
        self.code = code
        self.category = category
        self.brand = brand
        self.model = model
        self.packaging = packaging


# Phase 1 Task Foundation：显式加载模型供 Flask-Migrate metadata 使用，并注册最小操作接口。
from task_models import OrderTask, TaskActivity
from reminders.routes import reminders_bp
from reminders.cli import reminders_cli
from reminders.engine import reconcile_order_tasks_for_pi

app.register_blueprint(reminders_bp)
app.cli.add_command(reminders_cli)


class TargetedReconcileError(RuntimeError):
    """Safe boundary error for an atomic PI mutation/reconcile transaction."""


def commit_pi_with_targeted_reconcile(pi):
    """Flush PI facts, reconcile only this PI, then commit atomically."""
    try:
        db.session.flush()
        actions = reconcile_order_tasks_for_pi(pi)
        db.session.commit()
        return actions
    except Exception as exc:
        db.session.rollback()
        app.logger.exception("Targeted reminder reconcile failed for PI id=%s", pi.id)
        raise TargetedReconcileError(
            "订单保存失败，Reminder 同步未完成；本次修改已回滚。"
        ) from exc


DOCUMENT_REQUIRED_FACT_FIELDS = (
    'coc_required',
    'coa_required',
    'original_bl_required',
    'obd_electronic_required',
    'insurance_original_required',
    'insurance_electronic_required',
    'original_documents_mail_required',
    'telex_release_required',
)


def update_document_required_facts(pi, form):
    """Update only explicitly submitted Phase 2A tri-state Required Facts."""
    mapping = {'': None, 'true': True, 'false': False}
    for field in DOCUMENT_REQUIRED_FACT_FIELDS:
        if field not in form:
            continue
        value = form.get(field, '')
        if value not in mapping:
            raise ValueError(f'Invalid required fact value for {field}.')
        setattr(pi, field, mapping[value])


def update_shipping_facts(pi, form):
    """Update explicitly submitted Phase 2B facts and keep legacy date consumers aligned."""
    if 'container_loading_at' in form:
        raw_loading_at = (form.get('container_loading_at') or '').strip()
        pi.container_loading_at = (
            datetime.strptime(raw_loading_at, '%Y-%m-%dT%H:%M')
            if raw_loading_at
            else None
        )
        if pi.container_loading_at is not None:
            pi.container_date = pi.container_loading_at.date()
    for field in ('driver_name', 'driver_phone', 'vehicle_number'):
        if field in form:
            setattr(pi, field, (form.get(field) or '').strip() or None)


PAYMENT_MONEY_FIELDS = (
    'advance_payment_percent',
    'advance_payment_amount',
    'advance_received_amount',
    'balance_payment_amount',
    'balance_received_amount',
)


def update_payment_facts(pi, form):
    """Update only explicitly submitted Phase 2C facts; blank means unknown."""
    from decimal import Decimal, InvalidOperation

    if 'currency' in form:
        currency = (form.get('currency') or '').strip().upper() or None
        if currency is not None and len(currency) > 10:
            raise ValueError('currency cannot exceed 10 characters.')
        pi.currency = currency
    for field in PAYMENT_MONEY_FIELDS:
        if field not in form:
            continue
        raw_value = (form.get(field) or '').strip()
        if not raw_value:
            setattr(pi, field, None)
            continue
        try:
            value = Decimal(raw_value)
        except InvalidOperation as exc:
            raise ValueError(f'Invalid decimal value for {field}.') from exc
        if not value.is_finite():
            raise ValueError(f'{field} must be a finite decimal value.')
        if value < 0:
            raise ValueError(f'{field} cannot be negative.')
        if field == 'advance_payment_percent' and value > 100:
            raise ValueError('advance_payment_percent cannot exceed 100.')
        setattr(pi, field, value)
    for field in ('advance_received_at', 'balance_received_at'):
        if field not in form:
            continue
        raw_value = (form.get(field) or '').strip()
        setattr(
            pi,
            field,
            datetime.strptime(raw_value, '%Y-%m-%dT%H:%M') if raw_value else None,
        )


def update_freight_facts(pi, form):
    """Delegate Phase 2D freight fact parsing and confirmation invalidation."""
    from reminders.freight import update_freight_facts_from_form

    return update_freight_facts_from_form(pi, form)

# -----------------------------
# 工具函数：生成唯一客户编码 C001、C002...
# -----------------------------

def generate_next_customer_code():
    last = Customer.query.order_by(Customer.id.desc()).first()
    if last and last.code and last.code[1:].isdigit():
        next_number = int(last.code[1:]) + 1
    else:
        next_number = 1
    return f"C{next_number:03d}"

#生成唯一的出口商编码 E001、E002...

def generate_next_exporter_code():
    last = Exporter.query.order_by(Exporter.id.desc()).first()
    if last and last.code and last.code[1:].isdigit():
        next_number = int(last.code[1:]) + 1
    else:
        next_number = 1
    return f"E{next_number:03d}"

#生成唯一的厂家编码 F001、F002...

def generate_next_factory_code():
    last = Factory.query.order_by(Factory.id.desc()).first()
    if last and last.code and last.code[1:].isdigit():
        next_number = int(last.code[1:]) + 1
    else:
        next_number = 1
    return f"F{next_number:03d}"

#生成唯一的货代编码 FF001、FF002...

def generate_next_freight_forwarder_code():
    last = FreightForwarder.query.order_by(FreightForwarder.id.desc()).first()
    if last and last.code and last.code[2:].isdigit():
        next_number = int(last.code[2:]) + 1
    else:
        next_number = 1
    return f"FF{next_number:03d}"

#生成唯一的产品代码

def generate_next_product_code():
    last = Product.query.order_by(Product.id.desc()).first()
    if last and last.code and last.code[1:].isdigit():
        next_number = int(last.code[1:]) + 1
    else:
        next_number = 1
    return f"P{next_number:03d}"

#



# -----------------------------
# 首页仪表盘
# -----------------------------

from datetime import datetime, timedelta
from sqlalchemy import func



@app.route('/')
@login_required
def index():
    from reminders.dashboard import get_dashboard_tasks
    from reminders.routes import ensure_task_csrf_token

    # 获取全部 PI
    status_order = ['新建', '待发运', '已发运', '已到港', '已完成']
    pis = PI.query.order_by(PI.pi_date.desc()).all()
    pis = sorted(
        pis,
        key=lambda pi: (status_order.index(pi.status) if pi.status in status_order else len(status_order), -pi.pi_date.toordinal() if pi.pi_date else 0)
    )
    unfinished_count = PI.query.filter(PI.status != '已完成').count()
    shipment_total = db.session.query(func.sum(PIItem.quantity)).join(PI).filter(PI.status == '待发运').scalar() or 0

    # ========== 汇率获取逻辑 ========== #
    try:
        resp = requests.get('https://open.er-api.com/v6/latest/USD', timeout=3)
        data = resp.json()
        cny_rate = data['rates']['CNY']
    except Exception:
        cny_rate = None
    # Dashboard consumes persisted Task data only. Full business-rule reconcile
    # remains an explicit CLI/service operation and is never run by this GET.
    task_dashboard = get_dashboard_tasks()
    task_csrf_token = ensure_task_csrf_token()
    today = datetime.today().date()

    # 统计本月订单数量，确保pi_date不为None再用extract
    month_order_count = PI.query.filter(
        PI.pi_date.isnot(None),
        extract('year', PI.pi_date) == today.year,
        extract('month', PI.pi_date) == today.month
    ).count()

    return render_template(
        'index.html',
        pis=pis,
        unfinished_count=unfinished_count,
        shipment_total=shipment_total,
        task_dashboard=task_dashboard,
        task_csrf_token=task_csrf_token,
        cny_rate=cny_rate,  # 新增
        month_order_count=month_order_count  # 新增
    )

# -----------------------------
# 客户管理页面（添加 + 列表）
# -----------------------------

@app.route('/customers', methods=['GET', 'POST'])
@login_required
def customers():
    if request.method == 'POST':
        name = request.form['name']
        address = request.form['address']
        country = request.form['country']
        tax_code = request.form.get('tax_code', '')
        contact_person = request.form.get('contact_person', '')
        phone = request.form.get('phone', '')
        email = request.form.get('email', '')

        code = generate_next_customer_code()

        new_customer = Customer(
            code=code,
            name=name,
            address=address,
            tax_code=tax_code,
            country=country,
            contact_person=contact_person,
            phone=phone,
            email=email
        )
        db.session.add(new_customer)
        db.session.commit()
        return redirect(url_for('customers'))

    # GET请求：显示客户列表
    all_customers = Customer.query.order_by(Customer.code.asc()).all()
    return render_template('customers.html', customers=all_customers)


# -----------------------------
# 删除客户
# -----------------------------

@app.route('/delete-customer/<int:customer_id>')
@login_required
def delete_customer(customer_id):
    customer = Customer.query.get_or_404(customer_id)
    db.session.delete(customer)
    db.session.commit()
    return redirect(url_for('customers'))


# -----------------------------
# 编辑客户
# -----------------------------

@app.route('/edit-customer/<int:customer_id>', methods=['GET', 'POST'])
@login_required
def edit_customer(customer_id):
    customer = Customer.query.get_or_404(customer_id)

    if request.method == 'POST':
        customer.name = request.form['name']
        customer.address = request.form['address']
        customer.country = request.form['country']
        customer.tax_code = request.form.get('tax_code', '')
        customer.contact_person = request.form.get('contact_person', '')
        customer.phone = request.form.get('phone', '')
        customer.email = request.form.get('email', '')

        db.session.commit()
        return redirect(url_for('customers'))

    return render_template('edit_customer.html', customer=customer)
#
#出口商管理页面（列表+添加）
#
@app.route('/exporters', methods=['GET', 'POST'])
@login_required
def exporters():
    if request.method == 'POST':
        name = request.form['name']
        address = request.form['address']
        country = request.form['country']
        tax_code = request.form.get('tax_code', '')
        contact_person = request.form.get('contact_person', '')
        phone = request.form.get('phone', '')
        email = request.form.get('email', '')

        code = generate_next_exporter_code()

        new_exporter = Exporter(
            code=code,
            name=name,
            address=address,
            tax_code=tax_code,
            country=country,
            contact_person=contact_person,
            phone=phone,
            email=email
        )
        db.session.add(new_exporter)
        db.session.commit()
        return redirect(url_for('exporters'))

    all_exporters = Exporter.query.order_by(Exporter.code.asc()).all()
    return render_template('exporters.html', exporters=all_exporters)
#
#删除出口商
#
@app.route('/delete-exporter/<int:exporter_id>')
@login_required
def delete_exporter(exporter_id):
    exporter = Exporter.query.get_or_404(exporter_id)
    db.session.delete(exporter)
    db.session.commit()
    return redirect(url_for('exporters'))
#
#编辑出口商
#
@app.route('/edit-exporter/<int:exporter_id>', methods=['GET', 'POST'])
@login_required
def edit_exporter(exporter_id):
    exporter = Exporter.query.get_or_404(exporter_id)

    if request.method == 'POST':
        exporter.name = request.form['name']
        exporter.address = request.form['address']
        exporter.country = request.form['country']
        exporter.tax_code = request.form.get('tax_code', '')
        exporter.contact_person = request.form.get('contact_person', '')
        exporter.phone = request.form.get('phone', '')
        exporter.email = request.form.get('email', '')

        db.session.commit()
        return redirect(url_for('exporters'))

    return render_template('edit_exporter.html', exporter=exporter)

# -----------------------------
# 厂家管理页面（添加 + 列表）
# -----------------------------
@app.route('/factories', methods=['GET', 'POST'])
@login_required
def factories():
    if request.method == 'POST':
        name = request.form['name']
        country = request.form['country']
        address = request.form.get('address', '')
        tax_code = request.form.get('tax_code', '')
        contact_person = request.form.get('contact_person', '')
        phone = request.form.get('phone', '')
        email = request.form.get('email', '')

        code = generate_next_factory_code()

        new_factory = Factory(
            code=code,
            name=name,
            address=address,
            tax_code=tax_code,
            country=country,
            contact_person=contact_person,
            phone=phone,
            email=email
        )
        db.session.add(new_factory)
        db.session.commit()
        return redirect(url_for('factories'))

    all_factories = Factory.query.order_by(Factory.code.asc()).all()
    return render_template('factories.html', factories=all_factories)

# -----------------------------
# 删除厂家
# -----------------------------
@app.route('/delete-factory/<int:factory_id>')
@login_required
def delete_factory(factory_id):
    factory = Factory.query.get_or_404(factory_id)
    db.session.delete(factory)
    db.session.commit()
    return redirect(url_for('factories'))
# -----------------------------
# 编辑厂家
# -----------------------------

@app.route('/edit-factory/<int:factory_id>', methods=['GET', 'POST'])
@login_required
def edit_factory(factory_id):
    factory = Factory.query.get_or_404(factory_id)

    if request.method == 'POST':
        factory.name = request.form['name']
        factory.country = request.form['country']
        factory.address = request.form.get('address', '')
        factory.tax_code = request.form.get('tax_code', '')
        factory.contact_person = request.form.get('contact_person', '')
        factory.phone = request.form.get('phone', '')
        factory.email = request.form.get('email', '')

        db.session.commit()
        return redirect(url_for('factories'))

    return render_template('edit_factory.html', factory=factory)

# -----------------------------
# 货代管理页面（添加 + 列表）
# -----------------------------
@app.route('/freight-forwarders', methods=['GET', 'POST'])
@login_required
def freight_forwarders():
    if request.method == 'POST':
        name = request.form['name']
        address = request.form['address']
        country = request.form['country']
        tax_code = request.form.get('tax_code', '')
        contact_person = request.form.get('contact_person', '')
        phone = request.form.get('phone', '')
        email = request.form.get('email', '')

        code = generate_next_freight_forwarder_code()

        new_freight_forwarder = FreightForwarder(
            code=code,
            name=name,
            address=address,
            tax_code=tax_code,
            country=country,
            contact_person=contact_person,
            phone=phone,
            email=email
        )
        db.session.add(new_freight_forwarder)
        db.session.commit()
        return redirect(url_for('freight_forwarders'))

    all_freight_forwarders = FreightForwarder.query.order_by(FreightForwarder.code.asc()).all()
    return render_template('freight_forwarders.html', freight_forwarders=all_freight_forwarders)

# -----------------------------
# 删除货代
# -----------------------------
@app.route('/delete-freight-forwarder/<int:freight_forwarder_id>')
@login_required
def delete_freight_forwarder(freight_forwarder_id):
    freight_forwarder = FreightForwarder.query.get_or_404(freight_forwarder_id)
    db.session.delete(freight_forwarder)
    db.session.commit()
    return redirect(url_for('freight_forwarders'))

# -----------------------------
# 编辑货代
# -----------------------------
@app.route('/edit-freight-forwarder/<int:freight_forwarder_id>', methods=['GET', 'POST'])
@login_required
def edit_freight_forwarder(freight_forwarder_id):
    freight_forwarder = FreightForwarder.query.get_or_404(freight_forwarder_id)

    if request.method == 'POST':
        freight_forwarder.name = request.form['name']
        freight_forwarder.address = request.form['address']
        freight_forwarder.country = request.form['country']
        freight_forwarder.tax_code = request.form.get('tax_code', '')
        freight_forwarder.contact_person = request.form.get('contact_person', '')
        freight_forwarder.phone = request.form.get('phone', '')
        freight_forwarder.email = request.form.get('email', '')

        db.session.commit()
        return redirect(url_for('freight_forwarders'))

    return render_template('edit_freight_forwarder.html', freight_forwarder=freight_forwarder)

# -----------------------------
# 货代报价管理页面
# -----------------------------
@app.route('/freight-quotes', methods=['GET', 'POST'])
@login_required
def freight_quotes():
    if request.method == 'POST':
        freight_forwarder_id = int(request.form['freight_forwarder_id'])
        shipping_company = request.form['shipping_company']
        departure_port = request.form['departure_port']
        destination_port = request.form['destination_port']
        route_type = request.form['route_type']
        price = float(request.form['price'])
        quote_date = datetime.strptime(request.form['quote_date'], '%Y-%m-%d').date()
        valid_until = datetime.strptime(request.form['valid_until'], '%Y-%m-%d').date()
        destination_14_days = request.form.get('destination_14_days', '')
        remarks = request.form.get('remarks', '')

        new_quote = FreightQuote(
            freight_forwarder_id=freight_forwarder_id,
            shipping_company=shipping_company,
            departure_port=departure_port,
            destination_port=destination_port,
            route_type=route_type,
            price=price,
            quote_date=quote_date,
            valid_until=valid_until,
            destination_14_days=destination_14_days,
            remarks=remarks
        )
        db.session.add(new_quote)
        db.session.commit()
        return redirect(url_for('freight_quotes'))

    # GET请求：显示报价列表（支持筛选和排序）
    query = FreightQuote.query.join(FreightForwarder)
    
    # 筛选条件
    freight_forwarder_id = request.args.get('freight_forwarder')
    if freight_forwarder_id and freight_forwarder_id.isdigit():
        query = query.filter(FreightQuote.freight_forwarder_id == int(freight_forwarder_id))
    
    shipping_company = request.args.get('shipping_company')
    if shipping_company:
        query = query.filter(FreightQuote.shipping_company.ilike(f'%{shipping_company}%'))
    
    departure_port = request.args.get('departure_port')
    if departure_port:
        query = query.filter(FreightQuote.departure_port.ilike(f'%{departure_port}%'))
    
    destination_port = request.args.get('destination_port')
    if destination_port:
        query = query.filter(FreightQuote.destination_port.ilike(f'%{destination_port}%'))
    
    valid_until_filter = request.args.get('valid_until')
    if valid_until_filter:
        today = datetime.today().date()
        if valid_until_filter == 'valid':
            query = query.filter(FreightQuote.valid_until >= today)
        elif valid_until_filter == 'expired':
            query = query.filter(FreightQuote.valid_until < today)
    
    # 排序
    sort_by = request.args.get('sort', 'quote_date')
    order = request.args.get('order', 'desc')
    
    if sort_by == 'price':
        if order == 'asc':
            query = query.order_by(FreightQuote.price.asc())
        else:
            query = query.order_by(FreightQuote.price.desc())
    else:
        if order == 'asc':
            query = query.order_by(FreightQuote.quote_date.asc())
        else:
            query = query.order_by(FreightQuote.quote_date.desc())
    
    all_quotes = query.all()
    freight_forwarders = FreightForwarder.query.all()
    return render_template('freight_quotes.html', quotes=all_quotes, freight_forwarders=freight_forwarders)

# -----------------------------
# 删除报价
# -----------------------------
@app.route('/delete-freight-quote/<int:quote_id>')
@login_required
def delete_freight_quote(quote_id):
    quote = FreightQuote.query.get_or_404(quote_id)
    db.session.delete(quote)
    db.session.commit()
    return redirect(url_for('freight_quotes'))

# -----------------------------
# 编辑报价
# -----------------------------
@app.route('/edit-freight-quote/<int:quote_id>', methods=['GET', 'POST'])
@login_required
def edit_freight_quote(quote_id):
    quote = FreightQuote.query.get_or_404(quote_id)

    if request.method == 'POST':
        quote.freight_forwarder_id = int(request.form['freight_forwarder_id'])
        quote.shipping_company = request.form['shipping_company']
        quote.departure_port = request.form['departure_port']
        quote.destination_port = request.form['destination_port']
        quote.route_type = request.form['route_type']
        quote.price = float(request.form['price'])
        quote.quote_date = datetime.strptime(request.form['quote_date'], '%Y-%m-%d').date()
        quote.valid_until = datetime.strptime(request.form['valid_until'], '%Y-%m-%d').date()
        quote.destination_14_days = request.form.get('destination_14_days', '')
        quote.remarks = request.form.get('remarks', '')

        db.session.commit()
        return redirect(url_for('freight_quotes'))

    freight_forwarders = FreightForwarder.query.all()
    return render_template('edit_freight_quote.html', quote=quote, freight_forwarders=freight_forwarders)

# -----------------------------
# 产品管理页面（添加 + 列表）
# -----------------------------

@app.route('/products', methods=['GET', 'POST'])
@login_required
def products():
    if request.method == 'POST':
        category = request.form.get('category', 'Titanium dioxide')
        brand = request.form.get('brand', '')
        model = request.form['model']
        packaging = request.form.get('packaging', '')

        code = generate_next_product_code()

        new_product = Product(
            code=code,
            category=category,
            brand=brand,
            model=model,
            packaging=packaging
        )
        db.session.add(new_product)
        db.session.commit()
        return redirect(url_for('products'))

    all_products = Product.query.order_by(Product.code.asc()).all()
    return render_template('products.html', products=all_products)

#产品编辑

@app.route('/edit-product/<int:product_id>', methods=['GET', 'POST'])
@login_required
def edit_product(product_id):
    product = Product.query.get_or_404(product_id)

    if request.method == 'POST':
        product.category = request.form.get('category', 'Titanium dioxide')
        product.brand = request.form.get('brand', '')
        product.model = request.form['model']
        product.packaging = request.form.get('packaging', '')

        db.session.commit()
        return redirect(url_for('products'))

    return render_template('edit_product.html', product=product)

#产品删除

@app.route('/delete-product/<int:product_id>')
@login_required
def delete_product(product_id):
    product = Product.query.get_or_404(product_id)
    db.session.delete(product)
    db.session.commit()
    return redirect(url_for('products'))

# 检查PI编号是否重复的API接口
@app.route('/check-pi-no', methods=['POST'])
@login_required
def check_pi_no():
    pi_no = request.form.get('pi_no', '').strip()
    if not pi_no:
        return jsonify({'valid': False, 'message': 'PI编号不能为空'})
    
    # 检查是否已存在
    existing_pi = PI.query.filter_by(pi_no=pi_no).first()
    if existing_pi:
        return jsonify({'valid': False, 'message': f'PI编号 "{pi_no}" 已存在，请使用其他编号'})
    
    return jsonify({'valid': True, 'message': 'PI编号可用'})

# 创建PI

@app.route('/create-pi', methods=['GET', 'POST'])
@login_required
def create_pi():
    customers = Customer.query.all()
    exporters = Exporter.query.all()
    products = Product.query.all()
    factories = Factory.query.all()
    banks = ["Hang Seng Bank Limited", "BOC CHINA", "OCBC Bank HK"]

    if request.method == 'POST':
        # 获取表单基础字段
        pi_no = request.form['pi_no']
        pi_date = datetime.strptime(request.form['pi_date'], '%Y-%m-%d').date()
        customer_id = int(request.form['customer'])
        exporter_id = int(request.form['exporter'])
        payment_terms = request.form.get('payment_terms', '')
        loading_port = request.form.get('loading_port', '')
        destination_port = request.form.get('destination_port', '')
        bank = request.form.get('bank', '')
        shipment_date_str = request.form.get('shipment_date')
        shipment_date = datetime.strptime(shipment_date_str, '%Y-%m-%d').date() if shipment_date_str else None
        note = request.form.get('note', 'Please arrange shipment as scheduled.')

        # 校验 PI 编号不重复
        existing = PI.query.filter_by(pi_no=pi_no).first()
        if existing:
            return "⚠️ PI 编号已存在，请返回修改", 400

        # 获取客户和出口商信息用于快照
        customer = Customer.query.get(customer_id)
        exporter = Exporter.query.get(exporter_id)
        
        if customer is None or exporter is None:
            return "客户或出口商不存在", 400
        
        # 创建主表 PI 记录
        pi = PI(
            pi_no=pi_no,
            pi_date=pi_date,
            customer_id=customer_id,
            exporter_id=exporter_id,
            payment_terms=payment_terms,
            loading_port=loading_port,
            destination_port=destination_port,
            bank=bank,
            shipment_date=shipment_date,
            note=note,
            # 填充客户快照数据
            customer_name_snapshot=customer.name,
            customer_address_snapshot=customer.address,
            customer_tax_code_snapshot=customer.tax_code,
            customer_country_snapshot=customer.country,
            customer_contact_snapshot=customer.contact_person,
            customer_phone_snapshot=customer.phone,
            customer_email_snapshot=customer.email,
            # 填充出口商快照数据
            exporter_name_snapshot=exporter.name,
            exporter_address_snapshot=exporter.address,
            exporter_tax_code_snapshot=exporter.tax_code,
            exporter_country_snapshot=exporter.country,
            exporter_contact_snapshot=exporter.contact_person,
            exporter_phone_snapshot=exporter.phone,
            exporter_email_snapshot=exporter.email
        )
        db.session.add(pi)
        update_document_required_facts(pi, request.form)
        update_shipping_facts(pi, request.form)
        update_payment_facts(pi, request.form)
        update_freight_facts(pi, request.form)
        db.session.flush()  # 获取 pi.id

        # 创建产品明细行
        index = 0
        while f'product_{index}' in request.form:
            product_id = int(request.form[f'product_{index}'])
            factory_id = int(request.form[f'factory_{index}'])
            trade_term = request.form.get(f'trade_term_{index}', '')
            unit_price = float(request.form.get(f'unit_price_{index}', 0))
            quantity = float(request.form.get(f'quantity_{index}', 0))
            total_price = unit_price * quantity

            # 获取产品和厂家信息用于快照
            product = Product.query.get(product_id)
            factory = Factory.query.get(factory_id)
            
            if product is None or factory is None:
                return "产品或厂家不存在", 400

            item = PIItem(
                pi_id=pi.id,
                product_id=product_id,
                factory_id=factory_id,
                trade_term=trade_term,
                unit_price=unit_price,
                quantity=quantity,
                total_price=total_price,
                # 填充产品快照数据
                product_category_snapshot=product.category,
                product_brand_snapshot=product.brand,
                product_model_snapshot=product.model,
                product_packaging_snapshot=product.packaging,
                # 填充厂家快照数据
                factory_name_snapshot=factory.name,
                factory_address_snapshot=factory.address,
                factory_tax_code_snapshot=factory.tax_code,
                factory_country_snapshot=factory.country,
                factory_contact_snapshot=factory.contact_person,
                factory_phone_snapshot=factory.phone,
                factory_email_snapshot=factory.email
            )
            db.session.add(item)
            index += 1

        try:
            commit_pi_with_targeted_reconcile(pi)
        except TargetedReconcileError as exc:
            flash(str(exc), 'danger')
            return redirect(url_for('create_pi'))
        return redirect(url_for('index'))  # 成功后返回首页或可改为 pi_list 页面

    return render_template(
        'create_pi.html',
        customers=customers,
        exporters=exporters,
        products=products,
        factories=factories,
        banks=banks
    )

#
#PI list列表
#
@app.route('/pi-list')
@login_required
def show_pi_list():
    status_order = ['新建', '待发运', '已发运', '已到港', '已完成']
    # 只查询销售订单
    pi_list = PI.query.filter(PI.commission_factory_id.is_(None), PI.commission_exporter_id.is_(None)).all()
    # 按自定义状态顺序和PI日期倒序排序
    pi_list = sorted(
        pi_list,
        key=lambda pi: (status_order.index(pi.status) if pi.status in status_order else len(status_order), -pi.pi_date.toordinal() if pi.pi_date else 0)
    )
    
    # 准备图表数据
    chart_data = prepare_chart_data(pi_list)
    
    return render_template('pi_list.html', pi_list=pi_list, chart_data=chart_data)

def prepare_chart_data(pi_list):
    """准备图表数据"""
    from collections import defaultdict, Counter
    from datetime import datetime, timedelta
    
    # 1. 订单状态分布
    status_counts = Counter(pi.status for pi in pi_list)
    status_data = {
        'labels': list(status_counts.keys()),
        'data': list(status_counts.values()),
        'colors': ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0', '#9966FF']
    }
    
    # 2. 月度订单趋势（最近12个月）
    monthly_counts = defaultdict(int)
    current_date = datetime.now()
    for pi in pi_list:
        if pi.pi_date:
            month_key = pi.pi_date.strftime('%Y-%m')
            monthly_counts[month_key] += 1
    
    # 生成最近12个月的完整数据
    monthly_trend = []
    for i in range(11, -1, -1):
        month_date = current_date.replace(day=1) - timedelta(days=i*30)
        month_key = month_date.strftime('%Y-%m')
        monthly_trend.append({
            'month': month_date.strftime('%Y年%m月'),
            'count': monthly_counts.get(month_key, 0)
        })
    
    # 3. 客户订单金额排行（前10名）
    customer_amounts = defaultdict(float)
    for pi in pi_list:
        customer_name = pi.customer_name_snapshot or (pi.customer.name if pi.customer else '未知客户')
        total_amount = sum(item.total_price for item in pi.products) if pi.products else 0
        customer_amounts[customer_name] += total_amount
    
    top_customers = sorted(customer_amounts.items(), key=lambda x: x[1], reverse=True)[:10]
    customer_data = {
        'labels': [name for name, _ in top_customers],
        'data': [amount for _, amount in top_customers]
    }
    
    # 4. 产品销量统计
    product_quantities = defaultdict(float)
    for pi in pi_list:
        for item in pi.products:
            product_name = item.product_model_snapshot or (item.product.model if item.product else '未知产品')
            product_quantities[product_name] += item.quantity
    
    top_products = sorted(product_quantities.items(), key=lambda x: x[1], reverse=True)[:8]
    product_data = {
        'labels': [name for name, _ in top_products],
        'data': [qty for _, qty in top_products]
    }
    
    # 5. 总体统计
    total_stats = {
        'total_orders': len(pi_list),
        'total_amount': sum(sum(item.total_price for item in pi.products) for pi in pi_list if pi.products),
        'completed_orders': len([pi for pi in pi_list if pi.status == '已完成']),
        'pending_orders': len([pi for pi in pi_list if pi.status in ['新建', '待发运', '已发运', '已到港']])
    }
    
    return {
        'status_data': status_data,
        'monthly_trend': monthly_trend,
        'customer_data': customer_data,
        'product_data': product_data,
        'total_stats': total_stats
    }
#
#PI View查看
#
@app.route('/pi/<int:pi_id>')
@login_required
def view_pi(pi_id):
    pi = PI.query.get_or_404(pi_id)
    return render_template('view_pi.html', pi=pi)

#
#PI Edit编辑
#
@app.route('/pi/<int:pi_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_pi(pi_id):
    pi = PI.query.get_or_404(pi_id)
    
    # 添加调试信息
    print(f"🔍 PI状态检查: {pi.pi_no} 的状态是 '{pi.status}'")

    if request.method == 'POST':
        try:

            
            date_fmt = "%Y-%m-%d"
            pi.pi_no = request.form.get('pi_no')
            pi_date_str = request.form.get('pi_date')
            if pi_date_str:
                pi.pi_date = datetime.strptime(pi_date_str, date_fmt).date()
            new_customer_id = int(request.form['customer_id'])
            new_exporter_id = int(request.form['exporter_id'])
            
            # 强制更新字段，即使值相同
            pi.customer_id = new_customer_id
            pi.exporter_id = new_exporter_id
            
            pi.customer_id = new_customer_id
            pi.exporter_id = new_exporter_id
            pi.payment_terms = request.form.get('payment_terms')
            pi.loading_port = request.form.get('loading_port')
            pi.destination_port = request.form.get('destination_port')
            shipment_date_str = request.form.get('shipment_date')
            pi.shipment_date = datetime.strptime(shipment_date_str, date_fmt).date() if shipment_date_str else None
            pi.note = request.form.get('note')
            pi.bank = request.form.get('bank')
            pi.bill_of_lading = request.form.get('bill_of_lading', pi.bill_of_lading)
            pi.customs_docs_required = request.form.get('customs_docs_required', pi.customs_docs_required)
            pi.coo_required = request.form.get('coo_required', pi.coo_required)
            pi.apta_required = request.form.get('apta_required', pi.apta_required)
            pi.export_license_required = request.form.get('export_license_required', pi.export_license_required)
            update_document_required_facts(pi, request.form)
            update_payment_facts(pi, request.form)
            update_freight_facts(pi, request.form)
            pi.other_documents = request.form.get('other_documents', pi.other_documents)
            freight_forwarder_id_str = request.form.get('freight_forwarder_id')
            pi.freight_forwarder_id = int(freight_forwarder_id_str) if freight_forwarder_id_str else pi.freight_forwarder_id
            pi.ocean_freight = float(request.form.get('ocean_freight', pi.ocean_freight) or 0) if request.form.get('ocean_freight') is not None else pi.ocean_freight
            container_date_str = request.form.get('container_date')
            pi.container_date = datetime.strptime(container_date_str, date_fmt).date() if container_date_str else pi.container_date
            update_shipping_facts(pi, request.form)
            pi.container_location = request.form.get('container_location', pi.container_location)
            etd_str = request.form.get('etd')
            pi.etd = datetime.strptime(etd_str, date_fmt).date() if etd_str else pi.etd
            eta_str = request.form.get('eta')
            pi.eta = datetime.strptime(eta_str, date_fmt).date() if eta_str else pi.eta
            pi.vessel_info = request.form.get('vessel_info', pi.vessel_info)
            pi.booking_number = request.form.get('booking_number', pi.booking_number)
            pi.container_type_quantity = request.form.get('container_type_quantity', pi.container_type_quantity)
            pi.shipping_mark = request.form.get('shipping_mark', pi.shipping_mark)
            pi.freight_term = request.form.get('freight_term', pi.freight_term)
            pi.contract_number = request.form.get('contract_number', pi.contract_number)
            pi.freight_clause = request.form.get('freight_clause', pi.freight_clause)
            pi.waybill_option = request.form.get('waybill_option', pi.waybill_option)
            pi.container_number = request.form.get('container_number', pi.container_number)
            pi.seal_number = request.form.get('seal_number', pi.seal_number)
            pi.container_type = request.form.get('container_type', pi.container_type)
            pi.shipping_company = request.form.get('shipping_company', pi.shipping_company)
            actual_departure_date_str = request.form.get('actual_departure_date')
            pi.actual_departure_date = datetime.strptime(actual_departure_date_str, date_fmt).date() if actual_departure_date_str else pi.actual_departure_date
            pi.quantity_units = float(request.form.get('quantity_units', pi.quantity_units) or 0) if request.form.get('quantity_units') is not None else pi.quantity_units
            pi.gross_weight = float(request.form.get('gross_weight', pi.gross_weight) or 0) if request.form.get('gross_weight') is not None else pi.gross_weight
            pi.volume = float(request.form.get('volume', pi.volume) or 0) if request.form.get('volume') is not None else pi.volume
            pi.vgm = request.form.get('vgm', pi.vgm)
            pi.total_quantity = float(request.form.get('total_quantity', pi.total_quantity) or 0) if request.form.get('total_quantity') is not None else pi.total_quantity
            pi.total_quantity_unit = request.form.get('total_quantity_unit', pi.total_quantity_unit)
            pi.total_weight = float(request.form.get('total_weight', pi.total_weight) or 0) if request.form.get('total_weight') is not None else pi.total_weight
            pi.total_weight_unit = request.form.get('total_weight_unit', pi.total_weight_unit)
            pi.total_volume = float(request.form.get('total_volume', pi.total_volume) or 0) if request.form.get('total_volume') is not None else pi.total_volume
            pi.total_volume_unit = request.form.get('total_volume_unit', pi.total_volume_unit)
            pi.total_vgm = request.form.get('total_vgm', pi.total_vgm)
            # 已发运状态字段处理
            pi.bill_of_lading = request.form.get('bill_of_lading', pi.bill_of_lading)
            pi.shipping_company = request.form.get('shipping_company', pi.shipping_company)
            actual_departure_date_str = request.form.get('actual_departure_date')
            pi.actual_departure_date = datetime.strptime(actual_departure_date_str, date_fmt).date() if actual_departure_date_str else pi.actual_departure_date
            pi.batch_no = request.form.get('batch_no', pi.batch_no)
            pi.coa_status = request.form.get('coa_status', pi.coa_status)
            pi.insurance_status = request.form.get('insurance_status', pi.insurance_status)
            pi.document_shipping_status = request.form.get('document_shipping_status', pi.document_shipping_status)
            # 新增实际发运日期保存逻辑
            if pi.document_shipping_status == '已邮寄':
                pi.tracking_number = request.form.get('tracking_number')
            else:
                pi.tracking_number = None
            # 新增货款收齐状态保存逻辑
            pi.payment_received = request.form.get('payment_received')
            # 新增结汇文件需求保存逻辑
            pi.settlement_documents_required = request.form.get('settlement_documents_required')
            # 已到港状态字段处理
            actual_arrival_date_str = request.form.get('actual_arrival_date')
            pi.actual_arrival_date = datetime.strptime(actual_arrival_date_str, date_fmt).date() if actual_arrival_date_str else pi.actual_arrival_date
            pi.freight_invoice_amount = float(request.form.get('freight_invoice_amount', pi.freight_invoice_amount) or 0) if request.form.get('freight_invoice_amount') is not None else pi.freight_invoice_amount
            pi.freight_invoice_confirmed = request.form.get('freight_invoice_confirmed', pi.freight_invoice_confirmed)
            pi.freight_invoice_issued = request.form.get('freight_invoice_issued', pi.freight_invoice_issued)
            pi.telex_release = request.form.get('telex_release', pi.telex_release)
            # 已完成状态字段处理
            pi.freight_payment_status = request.form.get('freight_payment_status')
            
            # 处理产品明细的更新
            index = 0
            print(f"🔍 开始处理产品明细，表单字段: {[k for k in request.form.keys() if k.startswith('product_')]}")
            while f'product_{index}' in request.form:
                product_id = int(request.form[f'product_{index}'])
                factory_id = int(request.form[f'factory_{index}'])
                trade_term = request.form.get(f'trade_term_{index}', '')
                unit_price = float(request.form.get(f'unit_price_{index}', 0))
                quantity = float(request.form.get(f'quantity_{index}', 0))
                total_price = unit_price * quantity
                
                print(f"📦 处理产品明细 {index}: product_id={product_id}, factory_id={factory_id}, trade_term={trade_term}, unit_price={unit_price}, quantity={quantity}, total_price={total_price}")
                
                # 查找对应的产品明细项
                if index < len(pi.products):
                    # 更新现有的产品明细
                    item = pi.products[index]
                    
                    # 获取产品和厂家信息用于更新快照
                    product = Product.query.get(product_id)
                    factory = Factory.query.get(factory_id)
                    
                    if product is None or factory is None:
                        return "产品或厂家不存在", 400
                    
                    item.product_id = product_id
                    item.factory_id = factory_id
                    item.trade_term = trade_term
                    item.unit_price = unit_price
                    item.quantity = quantity
                    item.total_price = total_price
                    
                    # 更新快照数据
                    item.product_category_snapshot = product.category
                    item.product_brand_snapshot = product.brand
                    item.product_model_snapshot = product.model
                    item.product_packaging_snapshot = product.packaging
                    item.factory_name_snapshot = factory.name
                    item.factory_address_snapshot = factory.address
                    item.factory_tax_code_snapshot = factory.tax_code
                    item.factory_country_snapshot = factory.country
                    item.factory_contact_snapshot = factory.contact_person
                    item.factory_phone_snapshot = factory.phone
                    item.factory_email_snapshot = factory.email
                else:
                    # 添加新的产品明细
                    product = Product.query.get(product_id)
                    factory = Factory.query.get(factory_id)
                    
                    if product is None or factory is None:
                        return "产品或厂家不存在", 400
                    
                    new_item = PIItem(
                        pi_id=pi.id,
                        product_id=product_id,
                        factory_id=factory_id,
                        trade_term=trade_term,
                        unit_price=unit_price,
                        quantity=quantity,
                        total_price=total_price,
                        # 填充产品快照数据
                        product_category_snapshot=product.category,
                        product_brand_snapshot=product.brand,
                        product_model_snapshot=product.model,
                        product_packaging_snapshot=product.packaging,
                        # 填充厂家快照数据
                        factory_name_snapshot=factory.name,
                        factory_address_snapshot=factory.address,
                        factory_tax_code_snapshot=factory.tax_code,
                        factory_country_snapshot=factory.country,
                        factory_contact_snapshot=factory.contact_person,
                        factory_phone_snapshot=factory.phone,
                        factory_email_snapshot=factory.email
                    )
                    db.session.add(new_item)
                
                index += 1
            
            # 删除多余的产品明细（如果用户删除了某些行）
            while index < len(pi.products):
                db.session.delete(pi.products[index])
            
            # 添加调试信息 - 提交前检查
            print(f"🔍 提交前检查:")
            print(f"   pi.customer_id: {pi.customer_id} (类型: {type(pi.customer_id)})")
            print(f"   pi.exporter_id: {pi.exporter_id} (类型: {type(pi.exporter_id)})")
            print(f"   pi 对象是否被标记为已修改: {pi in db.session.dirty}")
            
            # PI 事实与该订单的 Reminder/Activity 在同一事务提交。
            commit_pi_with_targeted_reconcile(pi)
            
            # 添加调试信息 - 提交后重新查询
            pi_after_commit = PI.query.get(pi.id)
            print(f"💾 数据库提交后的值:")
            print(f"   customer_id: {pi_after_commit.customer_id} (类型: {type(pi_after_commit.customer_id)})")
            print(f"   exporter_id: {pi_after_commit.exporter_id} (类型: {type(pi_after_commit.exporter_id)})")
            
            return render_template('success_and_redirect.html', pi_id=pi.id)
        except Exception as e:
            print(f"❌ 保存过程中发生异常: {e}")
            print(f"❌ 异常类型: {type(e)}")
            import traceback
            print(f"❌ 异常堆栈: {traceback.format_exc()}")
            db.session.rollback()
            flash('保存失败：请检查输入内容或稍后重试。', 'danger')
            # 继续渲染编辑页面并保留输入内容
    customers = Customer.query.all()
    exporters = Exporter.query.all()
    products = Product.query.all()
    factories = Factory.query.all()
    freight_forwarders = FreightForwarder.query.all()
    return render_template('edit_pi.html', pi=pi, customers=customers, exporters=exporters, products=products, factories=factories, freight_forwarders=freight_forwarders)



#
#PI状态管理
#

@app.route('/pi/<int:pi_id>/update-status', methods=['GET', 'POST'])
@login_required
def update_pi_status(pi_id):
    pi = PI.query.get_or_404(pi_id)

    all_status = ['新建', '待发运', '已发运', '已到港', '已完成']
    current_index = all_status.index(pi.status)
    next_status_options = all_status[current_index + 1:] if current_index < len(all_status) - 1 else []

    if request.method == 'POST':
        new_status = request.form.get('new_status')
        remarks = request.form.get('remarks')

        if new_status and new_status in next_status_options:
            pi.status = new_status
            update_document_required_facts(pi, request.form)
            update_payment_facts(pi, request.form)
            update_freight_facts(pi, request.form)

            # 根据新状态处理对应字段
            date_fmt = "%Y-%m-%d"
            if new_status == '待发运':
                # 发运准备信息
                freight_forwarder_id = request.form.get('freight_forwarder_id')
                pi.freight_forwarder_id = int(freight_forwarder_id) if freight_forwarder_id else None
                pi.ocean_freight = request.form.get('ocean_freight') or None
                container_date_str = request.form.get('container_date')
                pi.container_date = datetime.strptime(container_date_str, date_fmt).date() if container_date_str else None
                update_shipping_facts(pi, request.form)
                pi.container_location = request.form.get('container_location')
                etd_str = request.form.get('etd')
                pi.etd = datetime.strptime(etd_str, date_fmt).date() if etd_str else pi.etd
                eta_str = request.form.get('eta')
                pi.eta = datetime.strptime(eta_str, date_fmt).date() if eta_str else pi.eta
                pi.coo_required = request.form.get('coo_required')
                pi.apta_required = request.form.get('apta_required')
                pi.export_license_required = request.form.get('export_license_required')
                pi.customs_docs_required = request.form.get('customs_docs_required')
                pi.other_documents = request.form.get('other_documents') or None

                # 托书信息
                pi.vessel_info = request.form.get('vessel_info')
                pi.booking_number = request.form.get('booking_number')
                pi.container_type_quantity = request.form.get('container_type_quantity')
                pi.shipping_mark = request.form.get('shipping_mark')
                pi.freight_term = request.form.get('freight_term')
                pi.contract_number = request.form.get('contract_number')
                pi.freight_clause = request.form.get('freight_clause')
                pi.waybill_option = request.form.get('waybill_option')
                pi.container_number = request.form.get('container_number')
                pi.seal_number = request.form.get('seal_number')
                pi.container_type = request.form.get('container_type')

                pi.quantity_units = float(request.form.get('quantity_units') or 0)
                pi.gross_weight = float(request.form.get('gross_weight') or 0)
                pi.volume = float(request.form.get('volume') or 0)
                pi.vgm = request.form.get('vgm')

                pi.total_quantity = float(request.form.get('total_quantity') or 0)
                pi.total_quantity_unit = request.form.get('total_quantity_unit')
                pi.total_weight = float(request.form.get('total_weight') or 0)
                pi.total_weight_unit = request.form.get('total_weight_unit')
                pi.total_volume = float(request.form.get('total_volume') or 0)
                pi.total_volume_unit = request.form.get('total_volume_unit')
                pi.total_vgm = request.form.get('total_vgm')

            # 代发与→已发运填写字段
            etd_str = request.form.get('etd')
            pi.etd = datetime.strptime(etd_str, date_fmt).date() if etd_str else pi.etd
            eta_str = request.form.get('eta')
            pi.eta = datetime.strptime(eta_str, date_fmt).date() if eta_str else pi.eta

        if new_status == '已发运':
                print("进入已发运表单逻辑")
                print("🟢 已进入已发运处理逻辑")
                print("✅ 提单号为:", request.form.get('bill_of_lading'))
                print("✅ 船公司为:", request.form.get('shipping_company'))
                print("✅ COA状态:", request.form.get('coa_status'))
                print("✅ 原件邮寄状态:", request.form.get('document_shipping_status'))
                pi.bill_of_lading = request.form.get('bill_of_lading')
                pi.shipping_company = request.form.get('shipping_company') or None
                pi.batch_no = request.form.get('batch_no') or None
                pi.coa_status = request.form.get('coa_status')
                pi.insurance_status = request.form.get('insurance_status')
                pi.document_shipping_status = request.form.get('document_shipping_status')
                # 新增实际发运日期保存逻辑
                actual_departure_date_str = request.form.get('actual_departure_date')
                pi.actual_departure_date = datetime.strptime(actual_departure_date_str, date_fmt).date() if actual_departure_date_str else pi.actual_departure_date
                if pi.document_shipping_status == '已邮寄':
                   pi.tracking_number = request.form.get('tracking_number') 
                else:
                   pi.tracking_number = None
                # 新增货款收齐状态保存逻辑
                pi.payment_received = request.form.get('payment_received')
                # 新增结汇文件需求保存逻辑
                pi.settlement_documents_required = request.form.get('settlement_documents_required') 

        # 已到港状态字段处理
        if new_status == '已到港':
            actual_arrival_date_str = request.form.get('actual_arrival_date')
            pi.actual_arrival_date = datetime.strptime(actual_arrival_date_str, date_fmt).date() if actual_arrival_date_str else pi.actual_arrival_date
            pi.payment_received = request.form.get('payment_received')
            pi.freight_invoice_amount = float(request.form.get('freight_invoice_amount') or 0) if request.form.get('freight_invoice_amount') else None
            pi.freight_invoice_confirmed = request.form.get('freight_invoice_confirmed')
            pi.freight_invoice_issued = request.form.get('freight_invoice_issued')
            pi.telex_release = request.form.get('telex_release')
        pi.freight_invoice_confirmed = request.form.get('freight_invoice_confirmed', pi.freight_invoice_confirmed)
        pi.freight_invoice_issued = request.form.get('freight_invoice_issued', pi.freight_invoice_issued)
        pi.telex_release = request.form.get('telex_release', pi.telex_release)

        # 已完成状态字段处理
        if new_status == '已完成':
            pi.payment_received = request.form.get('payment_received')
            pi.freight_payment_status = request.form.get('freight_payment_status')

        # 🚨 校验约束 - 只在状态更新为"已完成"时验证
        if new_status == '已完成':
            if pi.payment_received != '已收齐' or pi.freight_payment_status != '已付款':
                return "无法完成状态更新：需确认货款收齐且货代运费已付款", 400
        
        try:
            commit_pi_with_targeted_reconcile(pi)
        except TargetedReconcileError as exc:
            flash(str(exc), 'danger')
            return redirect(url_for('update_pi_status', pi_id=pi.id))
        return redirect(url_for('show_pi_list'))
    

    # ✅ GET 请求只渲染页面，不使用 POST 中的变量
    freight_forwarders = FreightForwarder.query.all()
    return render_template(
        'update_status.html',
        pi=pi,
        current_status=pi.status,
        next_status_options=next_status_options,
        freight_forwarders=freight_forwarders
    )
#
#自动托书生成功能
#

@app.route('/generate-booking/<int:pi_id>')
@login_required
def generate_booking(pi_id):
    pi = PI.query.get_or_404(pi_id)
    customer = pi.customer
    exporter = pi.exporter

    # 模板路径
    template_path = 'templates/word/BN-Sample.docx'
    output_dir = 'static/booking_docs'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f'booking_{pi.pi_no}.docx')

    # 加载模板
    doc = Document(template_path)

    # 替换字段字典
    replacements = {
        "{{pi_no}}": pi.pi_no,
        "{{quantity}}": str(int(sum([item.quantity for item in pi.products]))),

        "{{vessel_info}}": pi.vessel_info or "",
        "{{booking_number}}": pi.booking_number or "",

        "{{exporter_name}}": (exporter.name if exporter else pi.exporter_name_snapshot) or "",
        "{{exporter_address}}": ((exporter.address if exporter else pi.exporter_address_snapshot) or "").replace(";", "\n"),
        "{{exporter_tax_code}}": ((exporter.tax_code if exporter else pi.exporter_tax_code_snapshot) or "").replace(";", "\n"),

        "{{customer_name}}": customer.name or "",
        "{{customer_address}}": (customer.address or "").replace(";", "\n"),
        "{{customer_tax_code}}": (customer.tax_code or "").replace(";", "\n"),

        "{{loading_port}}": pi.loading_port or "",
        "{{destination_port}}": pi.destination_port or "",

        "{{container_type_quantity}}": pi.container_type_quantity or "",
        "{{shipping_mark}}": pi.shipping_mark or "",

        "{{product_category}}": "Titanium dioxide",
        "{{product_brand}}": pi.products[0].product.brand if pi.products else "",
        "{{product_model}}": pi.products[0].product.model if pi.products else "",

        "{{freight_term}}": pi.freight_term or "",
        "{{contract_number}}": pi.contract_number or "",
        "{{freight_clause}}": pi.freight_clause or "",
        "{{waybill_option}}": pi.waybill_option or "",
        "{{container_number}}": pi.container_number or "",
        "{{seal_number}}": pi.seal_number or "",
        "{{container_type}}": pi.container_type or "",

        "{{quantity_units}}": str(int(pi.quantity_units or 0)),
        "{{gross_weight}}": str(int(pi.gross_weight or 0)),
        "{{volume}}": str(int(pi.volume or 0)),
        "{{vgm}}": pi.vgm or "",

        "{{total_quantity}}": str(int(pi.total_quantity or 0)),
        "{{total_quantity_unit}}": pi.total_quantity_unit or "",
        "{{total_weight}}": str(int(pi.total_weight or 0)),
        "{{total_weight_unit}}": pi.total_weight_unit or "",
        "{{total_volume_unit}}": pi.total_volume_unit or "",
        "{{total_volume}}": str(int(pi.total_volume or 0)),
        "{{total_vgm}}": pi.total_vgm or ""
    }

    # ✅ 替换段落中的占位符（解决 run 拆段问题）
    def replace_paragraph_text(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        for key, val in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ''

    # 替换正文段落
    for paragraph in doc.paragraphs:
        replace_paragraph_text(paragraph)

    # 替换表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph)

    # 保存生成文档
    doc.save(output_path)

    # 下载文件
    return send_file(output_path, as_attachment=True)

#
#自动proforma invoice生成功能
#
from flask import render_template
from weasyprint import HTML

@app.route('/generate-invoice/<int:pi_id>')
@login_required
def generate_invoice_pdf(pi_id):
    pi = PI.query.get_or_404(pi_id)
    rendered = render_template('proforma_invoice.html', pi=pi)
    pdf_path = f'static/invoices/proforma_invoice_{pi.pi_no}.pdf'
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    HTML(string=rendered).write_pdf(pdf_path)
    return send_file(pdf_path, as_attachment=True)

#
#自动Packing List生成功能
#
@app.route('/generate-packing-list/<int:pi_id>')
@login_required
def generate_packing_list_pdf(pi_id):
    pi = PI.query.get_or_404(pi_id)
    rendered = render_template('packing_list.html', pi=pi)
    pdf_path = f'static/packing_lists/packing_list_{pi.pi_no}.pdf'
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    HTML(string=rendered).write_pdf(pdf_path)
    return send_file(pdf_path, as_attachment=True)

#
#自动invoice生成功能
#
@app.route('/generate-invoice-pdf/<int:pi_id>')
@login_required
def generate_invoice_pdf_file(pi_id):
    pi = PI.query.get_or_404(pi_id)
    rendered = render_template('invoice.html', pi=pi)
    pdf_path = f'static/invoices/invoice_{pi.pi_no}.pdf'
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    HTML(string=rendered).write_pdf(pdf_path)
    return send_file(pdf_path, as_attachment=True)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('已安全退出', 'success')
    return redirect(url_for('login'))

# 创建工厂直发订单PI
@app.route('/create-pi-commission', methods=['GET', 'POST'])
@login_required
def create_pi_commission():
    customers = Customer.query.all()
    exporters = Exporter.query.all()
    products = Product.query.all()
    factories = Factory.query.all()
    banks = ["Hang Seng Bank Limited", "BOC CHINA", "OCBC Bank HK"]

    if request.method == 'POST':
        pi_no = request.form['pi_no']
        pi_date = datetime.strptime(request.form['pi_date'], '%Y-%m-%d').date()
        customer_id = int(request.form['customer'])
        payment_terms = request.form.get('payment_terms', '')
        loading_port = request.form.get('loading_port', '')
        destination_port = request.form.get('destination_port', '')
        bank = request.form.get('bank', '')
        shipment_date_str = request.form.get('shipment_date')
        shipment_date = datetime.strptime(shipment_date_str, '%Y-%m-%d').date() if shipment_date_str else None
        note = request.form.get('note', 'Please arrange shipment as scheduled.')
        # 工厂直发订单专属字段
        commission_factory_id = int(request.form['commission_factory_id'])
        commission_exporter_id = int(request.form['commission_exporter_id'])
        commission_amount = request.form.get('commission_amount')
        commission_rate = request.form.get('commission_rate')
        commission_status = request.form.get('commission_status')
        factory_sale_amount = request.form.get('factory_sale_amount')

        # 校验 PI 编号不重复
        existing = PI.query.filter_by(pi_no=pi_no).first()
        if existing:
            return "⚠️ PI 编号已存在，请返回修改", 400

        # 获取客户、厂家、佣金收取方信息用于快照
        customer = Customer.query.get(customer_id)
        factory = Factory.query.get(commission_factory_id)
        exporter = Exporter.query.get(commission_exporter_id)
        if customer is None or factory is None or exporter is None:
            return "客户、厂家或佣金收取方不存在", 400

        # 创建主表 PI 记录
        pi = PI(
            pi_no=pi_no,
            pi_date=pi_date,
            customer_id=customer_id,
            payment_terms=payment_terms,
            loading_port=loading_port,
            destination_port=destination_port,
            bank=bank,
            shipment_date=shipment_date,
            note=note,
            # 工厂直发订单专属字段
            commission_factory_id=commission_factory_id,
            commission_exporter_id=commission_exporter_id,
            commission_amount=float(commission_amount) if commission_amount else None,
            commission_rate=float(commission_rate) if commission_rate else None,
            commission_status=commission_status,
            factory_sale_amount=float(factory_sale_amount) if factory_sale_amount else None,
            # 快照
            customer_name_snapshot=customer.name,
            customer_address_snapshot=customer.address,
            customer_tax_code_snapshot=customer.tax_code,
            customer_country_snapshot=customer.country,
            customer_contact_snapshot=customer.contact_person,
            customer_phone_snapshot=customer.phone,
            customer_email_snapshot=customer.email,
            exporter_name_snapshot=exporter.name,
            exporter_address_snapshot=exporter.address,
            exporter_tax_code_snapshot=exporter.tax_code,
            exporter_country_snapshot=exporter.country,
            exporter_contact_snapshot=exporter.contact_person,
            exporter_phone_snapshot=exporter.phone,
            exporter_email_snapshot=exporter.email
        )
        db.session.add(pi)
        update_document_required_facts(pi, request.form)
        update_shipping_facts(pi, request.form)
        update_payment_facts(pi, request.form)
        update_freight_facts(pi, request.form)
        db.session.flush()  # 获取 pi.id

        # 创建产品明细行（如有产品明细表单，可参考销售订单逻辑）
        index = 0
        while f'product_{index}' in request.form:
            product_id = int(request.form[f'product_{index}'])
            factory_id = int(request.form[f'factory_{index}'])
            trade_term = request.form.get(f'trade_term_{index}', '')
            unit_price = float(request.form.get(f'unit_price_{index}', 0))
            quantity = float(request.form.get(f'quantity_{index}', 0))
            total_price = unit_price * quantity

            product = Product.query.get(product_id)
            factory = Factory.query.get(factory_id)
            if product is None or factory is None:
                return "产品或厂家不存在", 400

            item = PIItem(
                pi_id=pi.id,
                product_id=product_id,
                factory_id=factory_id,
                trade_term=trade_term,
                unit_price=unit_price,
                quantity=quantity,
                total_price=total_price,
                product_category_snapshot=product.category,
                product_brand_snapshot=product.brand,
                product_model_snapshot=product.model,
                product_packaging_snapshot=product.packaging,
                factory_name_snapshot=factory.name,
                factory_address_snapshot=factory.address,
                factory_tax_code_snapshot=factory.tax_code,
                factory_country_snapshot=factory.country,
                factory_contact_snapshot=factory.contact_person,
                factory_phone_snapshot=factory.phone,
                factory_email_snapshot=factory.email
            )
            db.session.add(item)
            index += 1

        try:
            commit_pi_with_targeted_reconcile(pi)
        except TargetedReconcileError as exc:
            flash(str(exc), 'danger')
            return redirect(url_for('create_pi_commission'))
        return redirect(url_for('index'))

    return render_template('create_pi_commission.html', customers=customers, exporters=exporters, factories=factories, products=products, banks=banks)

@app.route('/commission-pi-list')
@login_required
def commission_pi_list():
    # 只查询工厂直发订单，并预加载产品信息
    pi_list = PI.query.filter((PI.commission_factory_id.isnot(None)) | (PI.commission_exporter_id.isnot(None))).options(
        db.joinedload(PI.products).joinedload(PIItem.product)
    ).all()
    
    # 准备图表数据
    chart_data = prepare_commission_chart_data(pi_list)
    
    return render_template('commission_pi_list.html', pi_list=pi_list, chart_data=chart_data)

def prepare_commission_chart_data(pi_list):
    """准备工厂直发订单图表数据"""
    from collections import defaultdict, Counter
    from datetime import datetime, timedelta
    
    # 1. 状态分布统计
    status_counts = Counter()
    for pi in pi_list:
        status_counts[pi.status] += 1
    
    # 2. 月度趋势统计
    monthly_counts = defaultdict(int)
    for pi in pi_list:
        if pi.pi_date:
            month_key = pi.pi_date.strftime('%Y-%m')
            monthly_counts[month_key] += 1
    
    # 3. 客户订单金额排行
    customer_amounts = defaultdict(float)
    for pi in pi_list:
        customer_name = pi.customer_name_snapshot or (pi.customer.name if pi.customer else '未知客户')
        total_amount = sum(item.total_price for item in pi.products) if pi.products else 0
        customer_amounts[customer_name] += total_amount
    
    # 按金额排序，取前8名
    top_customers = sorted(customer_amounts.items(), key=lambda x: x[1], reverse=True)[:8]
    
    # 4. 厂家订单金额排行
    factory_amounts = defaultdict(float)
    for pi in pi_list:
        factory_name = pi.commission_factory.name if pi.commission_factory else '未知厂家'
        total_amount = sum(item.total_price for item in pi.products) if pi.products else 0
        factory_amounts[factory_name] += total_amount
    
    # 按金额排序，取前8名
    top_factories = sorted(factory_amounts.items(), key=lambda x: x[1], reverse=True)[:8]
    
    # 5. 产品销量统计
    product_quantities = defaultdict(float)
    for pi in pi_list:
        for item in pi.products:
            product_name = item.product_model_snapshot or (item.product.model if item.product else '未知产品')
            product_quantities[product_name] += item.quantity
    
    # 按销量排序，取前8名
    top_products = sorted(product_quantities.items(), key=lambda x: x[1], reverse=True)[:8]
    

    
    # 7. 总体统计
    total_orders = len(pi_list)
    completed_orders = sum(1 for pi in pi_list if pi.status == '已完成')
    in_progress_orders = total_orders - completed_orders
    total_amount = sum(sum(item.total_price for item in pi.products) for pi in pi_list if pi.products)
    total_commission = sum(pi.commission_amount for pi in pi_list if pi.commission_amount)
    
    return {
        'statusData': {
            'labels': list(status_counts.keys()),
            'data': list(status_counts.values())
        },
        'trendData': {
            'labels': sorted(monthly_counts.keys()),
            'data': [monthly_counts[month] for month in sorted(monthly_counts.keys())]
        },
        'customerData': {
            'labels': [customer[0] for customer in top_customers],
            'data': [customer[1] for customer in top_customers]
        },
        'factoryData': {
            'labels': [factory[0] for factory in top_factories],
            'data': [factory[1] for factory in top_factories]
        },
        'productData': {
            'labels': [product[0] for product in top_products],
            'data': [product[1] for product in top_products]
        },
        'totalStats': {
            'totalOrders': total_orders,
            'completedOrders': completed_orders,
            'inProgressOrders': in_progress_orders,
            'totalAmount': total_amount,
            'totalCommission': total_commission
        }
    }

@app.route('/commission-pi/<int:pi_id>/update-status', methods=['GET', 'POST'])
@login_required
def update_commission_pi_status(pi_id):
    pi = PI.query.get_or_404(pi_id)
    # 只允许工厂直发订单
    if not (pi.commission_factory_id or pi.commission_exporter_id):
        return "该订单不是工厂直发订单", 400
    all_status = ['新建', '待发运', '已发运', '已到港', '已完成']
    current_index = all_status.index(pi.status)
    next_status_options = all_status[current_index + 1:] if current_index < len(all_status) - 1 else []
    if request.method == 'POST':
        new_status = request.form.get('new_status')
        remarks = request.form.get('remarks')
        commission_status = request.form.get('commission_status', pi.commission_status)
        if new_status and new_status in next_status_options:
            pi.status = new_status
            update_document_required_facts(pi, request.form)
            update_payment_facts(pi, request.form)
            update_freight_facts(pi, request.form)
            
            # 根据新状态处理对应字段
            date_fmt = "%Y-%m-%d"
            
            # 待发运状态字段处理
            if new_status == '待发运':
                # 发运准备信息
                freight_forwarder_id = request.form.get('freight_forwarder_id')
                pi.freight_forwarder_id = int(freight_forwarder_id) if freight_forwarder_id else None
                pi.ocean_freight = request.form.get('ocean_freight') or None
                container_date_str = request.form.get('container_date')
                pi.container_date = datetime.strptime(container_date_str, date_fmt).date() if container_date_str else None
                update_shipping_facts(pi, request.form)
                pi.container_location = request.form.get('container_location')
                etd_str = request.form.get('etd')
                pi.etd = datetime.strptime(etd_str, date_fmt).date() if etd_str else pi.etd
                eta_str = request.form.get('eta')
                pi.eta = datetime.strptime(eta_str, date_fmt).date() if eta_str else pi.eta
                pi.coo_required = request.form.get('coo_required')
                pi.apta_required = request.form.get('apta_required')
                pi.export_license_required = request.form.get('export_license_required')
                pi.customs_docs_required = request.form.get('customs_docs_required')
                pi.other_documents = request.form.get('other_documents') or None

                # 托书信息
                pi.vessel_info = request.form.get('vessel_info')
                pi.booking_number = request.form.get('booking_number')
                pi.container_type_quantity = request.form.get('container_type_quantity')
                pi.shipping_mark = request.form.get('shipping_mark')
                pi.freight_term = request.form.get('freight_term')
                pi.contract_number = request.form.get('contract_number')
                pi.freight_clause = request.form.get('freight_clause')
                pi.waybill_option = request.form.get('waybill_option')
                pi.container_number = request.form.get('container_number')
                pi.seal_number = request.form.get('seal_number')
                pi.container_type = request.form.get('container_type')
                pi.quantity_units = float(request.form.get('quantity_units') or 0)
                pi.gross_weight = float(request.form.get('gross_weight') or 0)
                pi.volume = float(request.form.get('volume') or 0)
                pi.vgm = request.form.get('vgm')
                pi.total_quantity = float(request.form.get('total_quantity') or 0)
                pi.total_quantity_unit = request.form.get('total_quantity_unit')
                pi.total_weight = float(request.form.get('total_weight') or 0)
                pi.total_weight_unit = request.form.get('total_weight_unit')
                pi.total_volume = float(request.form.get('total_volume') or 0)
                pi.total_volume_unit = request.form.get('total_volume_unit')
                pi.total_vgm = request.form.get('total_vgm')

            # 已发运状态字段处理
            if new_status == '已发运':
                pi.bill_of_lading = request.form.get('bill_of_lading')
                pi.shipping_company = request.form.get('shipping_company') or None
                pi.batch_no = request.form.get('batch_no') or None
                pi.coa_status = request.form.get('coa_status')
                pi.insurance_status = request.form.get('insurance_status')
                pi.document_shipping_status = request.form.get('document_shipping_status')
                # 实际发运日期
                actual_departure_date_str = request.form.get('actual_departure_date')
                pi.actual_departure_date = datetime.strptime(actual_departure_date_str, date_fmt).date() if actual_departure_date_str else None
                if pi.document_shipping_status == '已邮寄':
                   pi.tracking_number = request.form.get('tracking_number') 
                else:
                   pi.tracking_number = None
                # 货款收齐状态
                pi.payment_received = request.form.get('payment_received')
                # 结汇文件需求
                pi.settlement_documents_required = request.form.get('settlement_documents_required')

            # 已到港状态字段处理
            if new_status == '已到港':
                actual_arrival_date_str = request.form.get('actual_arrival_date')
                pi.actual_arrival_date = datetime.strptime(actual_arrival_date_str, date_fmt).date() if actual_arrival_date_str else pi.actual_arrival_date
                pi.payment_received = request.form.get('payment_received')
                pi.freight_invoice_amount = float(request.form.get('freight_invoice_amount') or 0) if request.form.get('freight_invoice_amount') else None
                pi.freight_invoice_confirmed = request.form.get('freight_invoice_confirmed')
                pi.freight_invoice_issued = request.form.get('freight_invoice_issued')
                pi.telex_release = request.form.get('telex_release')

            # 已完成状态字段处理
            if new_status == '已完成':
                pi.payment_received = request.form.get('payment_received')
                pi.freight_payment_status = request.form.get('freight_payment_status')
                pi.commission_status = commission_status
                # 新增校验：佣金结算状态必须为已结算
                if pi.payment_received != '已收齐' or pi.freight_payment_status != '已付款' or pi.commission_status != '已结算':
                    return "无法完成状态更新：需确认货款收齐、货代运费已付款且佣金已结算", 400
            
            try:
                commit_pi_with_targeted_reconcile(pi)
            except TargetedReconcileError as exc:
                flash(str(exc), 'danger')
                return redirect(url_for('update_commission_pi_status', pi_id=pi.id))
            return redirect(url_for('commission_pi_list'))
    freight_forwarders = FreightForwarder.query.all()
    return render_template('update_commission_status.html', pi=pi, current_status=pi.status, next_status_options=next_status_options, freight_forwarders=freight_forwarders)

@app.route('/commission-pi/<int:pi_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_commission_pi(pi_id):
    pi = PI.query.get_or_404(pi_id)
    if not (pi.commission_factory_id or pi.commission_exporter_id):
        return "该订单不是工厂直发订单", 400
    customers = Customer.query.all()
    exporters = Exporter.query.all()
    products = Product.query.all()
    factories = Factory.query.all()
    freight_forwarders = FreightForwarder.query.all()
    banks = ["Hang Seng Bank Limited", "BOC CHINA", "OCBC Bank HK"]
    if request.method == 'POST':
        try:
            date_fmt = "%Y-%m-%d"
            # 基本信息字段
            pi.pi_no = request.form.get('pi_no')
            pi_date_str = request.form.get('pi_date')
            if pi_date_str:
                pi.pi_date = datetime.strptime(pi_date_str, date_fmt).date()
            pi.customer_id = int(request.form.get('customer_id'))
            pi.commission_factory_id = int(request.form.get('commission_factory_id'))
            pi.commission_exporter_id = int(request.form.get('commission_exporter_id'))
            pi.payment_terms = request.form.get('payment_terms')
            pi.loading_port = request.form.get('loading_port')
            pi.destination_port = request.form.get('destination_port')
            shipment_date_str = request.form.get('shipment_date')
            pi.shipment_date = datetime.strptime(shipment_date_str, date_fmt).date() if shipment_date_str else None
            pi.note = request.form.get('note')
            pi.bank = request.form.get('bank')
            pi.commission_amount = float(request.form.get('commission_amount') or 0)
            pi.commission_rate = float(request.form.get('commission_rate') or 0)
            pi.commission_status = request.form.get('commission_status')
            pi.factory_sale_amount = float(request.form.get('factory_sale_amount') or 0)
            
            # 发运准备信息字段
            if 'freight_forwarder_id' in request.form:
                freight_forwarder_id = request.form.get('freight_forwarder_id')
                pi.freight_forwarder_id = int(freight_forwarder_id) if freight_forwarder_id else None
            pi.ocean_freight = float(request.form.get('ocean_freight', pi.ocean_freight) or 0) if request.form.get('ocean_freight') is not None else pi.ocean_freight
            container_date_str = request.form.get('container_date')
            pi.container_date = datetime.strptime(container_date_str, date_fmt).date() if container_date_str else pi.container_date
            update_shipping_facts(pi, request.form)
            pi.container_location = request.form.get('container_location', pi.container_location)
            etd_str = request.form.get('etd')
            pi.etd = datetime.strptime(etd_str, date_fmt).date() if etd_str else pi.etd
            eta_str = request.form.get('eta')
            pi.eta = datetime.strptime(eta_str, date_fmt).date() if eta_str else pi.eta
            pi.coo_required = request.form.get('coo_required', pi.coo_required)
            pi.apta_required = request.form.get('apta_required', pi.apta_required)
            pi.export_license_required = request.form.get('export_license_required', pi.export_license_required)
            pi.customs_docs_required = request.form.get('customs_docs_required', pi.customs_docs_required)
            update_document_required_facts(pi, request.form)
            update_payment_facts(pi, request.form)
            update_freight_facts(pi, request.form)
            pi.other_documents = request.form.get('other_documents', pi.other_documents)
            
            # 托书信息字段
            pi.vessel_info = request.form.get('vessel_info', pi.vessel_info)
            pi.booking_number = request.form.get('booking_number', pi.booking_number)
            pi.container_type_quantity = request.form.get('container_type_quantity', pi.container_type_quantity)
            pi.shipping_mark = request.form.get('shipping_mark', pi.shipping_mark)
            pi.freight_term = request.form.get('freight_term', pi.freight_term)
            pi.contract_number = request.form.get('contract_number', pi.contract_number)
            pi.freight_clause = request.form.get('freight_clause', pi.freight_clause)
            pi.waybill_option = request.form.get('waybill_option', pi.waybill_option)
            pi.container_number = request.form.get('container_number', pi.container_number)
            pi.seal_number = request.form.get('seal_number', pi.seal_number)
            pi.container_type = request.form.get('container_type', pi.container_type)
            pi.quantity_units = float(request.form.get('quantity_units', pi.quantity_units) or 0) if request.form.get('quantity_units') is not None else pi.quantity_units
            pi.gross_weight = float(request.form.get('gross_weight', pi.gross_weight) or 0) if request.form.get('gross_weight') is not None else pi.gross_weight
            pi.volume = float(request.form.get('volume', pi.volume) or 0) if request.form.get('volume') is not None else pi.volume
            pi.vgm = request.form.get('vgm', pi.vgm)
            pi.total_quantity = float(request.form.get('total_quantity', pi.total_quantity) or 0) if request.form.get('total_quantity') is not None else pi.total_quantity
            pi.total_quantity_unit = request.form.get('total_quantity_unit', pi.total_quantity_unit)
            pi.total_weight = float(request.form.get('total_weight', pi.total_weight) or 0) if request.form.get('total_weight') is not None else pi.total_weight
            pi.total_weight_unit = request.form.get('total_weight_unit', pi.total_weight_unit)
            pi.total_volume = float(request.form.get('total_volume', pi.total_volume) or 0) if request.form.get('total_volume') is not None else pi.total_volume
            pi.total_volume_unit = request.form.get('total_volume_unit', pi.total_volume_unit)
            pi.total_vgm = request.form.get('total_vgm', pi.total_vgm)
            
            # 已发运信息字段
            pi.bill_of_lading = request.form.get('bill_of_lading', pi.bill_of_lading)
            pi.shipping_company = request.form.get('shipping_company', pi.shipping_company)
            actual_departure_date_str = request.form.get('actual_departure_date')
            pi.actual_departure_date = datetime.strptime(actual_departure_date_str, date_fmt).date() if actual_departure_date_str else pi.actual_departure_date
            pi.batch_no = request.form.get('batch_no', pi.batch_no)
            pi.coa_status = request.form.get('coa_status', pi.coa_status)
            pi.insurance_status = request.form.get('insurance_status', pi.insurance_status)
            pi.document_shipping_status = request.form.get('document_shipping_status', pi.document_shipping_status)
            # 新增实际发运日期保存逻辑
            if pi.document_shipping_status == '已邮寄':
                pi.tracking_number = request.form.get('tracking_number')
            else:
                pi.tracking_number = None
                
            # 已到港信息字段
            actual_arrival_date_str = request.form.get('actual_arrival_date')
            pi.actual_arrival_date = datetime.strptime(actual_arrival_date_str, date_fmt).date() if actual_arrival_date_str else pi.actual_arrival_date
            pi.payment_received = request.form.get('payment_received')
            pi.freight_invoice_amount = float(request.form.get('freight_invoice_amount', pi.freight_invoice_amount) or 0) if request.form.get('freight_invoice_amount') is not None else pi.freight_invoice_amount
            pi.freight_invoice_confirmed = request.form.get('freight_invoice_confirmed', pi.freight_invoice_confirmed)
            pi.freight_invoice_issued = request.form.get('freight_invoice_issued', pi.freight_invoice_issued)
            pi.telex_release = request.form.get('telex_release', pi.telex_release)
            pi.settlement_documents_required = request.form.get('settlement_documents_required')
            
            # 已完成状态字段处理
            pi.freight_payment_status = request.form.get('freight_payment_status')
            
            # 处理产品明细的更新
            index = 0
            print(f"🔍 开始处理产品明细，表单字段: {[k for k in request.form.keys() if k.startswith('product_')]}")
            while f'product_{index}' in request.form:
                product_id = int(request.form[f'product_{index}'])
                factory_id = int(request.form[f'factory_{index}'])
                trade_term = request.form.get(f'trade_term_{index}', '')
                unit_price = float(request.form.get(f'unit_price_{index}', 0))
                quantity = float(request.form.get(f'quantity_{index}', 0))
                total_price = unit_price * quantity
                
                if index < len(pi.products):
                    # 更新现有的产品明细
                    item = pi.products[index]
                    
                    # 获取产品和厂家信息用于更新快照
                    product = Product.query.get(product_id)
                    factory = Factory.query.get(factory_id)
                    
                    if product is None or factory is None:
                        return "产品或厂家不存在", 400
                    
                    item.product_id = product_id
                    item.factory_id = factory_id
                    item.trade_term = trade_term
                    item.unit_price = unit_price
                    item.quantity = quantity
                    item.total_price = total_price
                    
                    # 更新快照数据
                    item.product_category_snapshot = product.category
                    item.product_brand_snapshot = product.brand
                    item.product_model_snapshot = product.model
                    item.product_packaging_snapshot = product.packaging
                    item.factory_name_snapshot = factory.name
                    item.factory_address_snapshot = factory.address
                    item.factory_tax_code_snapshot = factory.tax_code
                    item.factory_country_snapshot = factory.country
                    item.factory_contact_snapshot = factory.contact_person
                    item.factory_phone_snapshot = factory.phone
                    item.factory_email_snapshot = factory.email
                else:
                    # 创建新的产品明细
                    product = Product.query.get(product_id)
                    factory = Factory.query.get(factory_id)
                    
                    if product is None or factory is None:
                        return "产品或厂家不存在", 400
                    
                    item = PIItem(
                        pi_id=pi.id,
                        product_id=product_id,
                        factory_id=factory_id,
                        trade_term=trade_term,
                        unit_price=unit_price,
                        quantity=quantity,
                        total_price=total_price,
                        product_category_snapshot=product.category,
                        product_brand_snapshot=product.brand,
                        product_model_snapshot=product.model,
                        product_packaging_snapshot=product.packaging,
                        factory_name_snapshot=factory.name,
                        factory_address_snapshot=factory.address,
                        factory_tax_code_snapshot=factory.tax_code,
                        factory_country_snapshot=factory.country,
                        factory_contact_snapshot=factory.contact_person,
                        factory_phone_snapshot=factory.phone,
                        factory_email_snapshot=factory.email
                    )
                    db.session.add(item)
                index += 1
            
            # 删除多余的产品明细
            while len(pi.products) > index:
                db.session.delete(pi.products[-1])
            
            commit_pi_with_targeted_reconcile(pi)
            return redirect(url_for('commission_pi_list'))
        except Exception as e:
            db.session.rollback()
            flash('保存失败：请检查输入内容或稍后重试。', 'danger')
            return redirect(url_for('edit_commission_pi', pi_id=pi.id))
    return render_template('edit_commission_pi.html', pi=pi, customers=customers, exporters=exporters, factories=factories, products=products, freight_forwarders=freight_forwarders, banks=banks)

# -----------------------------
# 启动应用
# -----------------------------

if __name__ == '__main__':
    app.run(debug=True)
