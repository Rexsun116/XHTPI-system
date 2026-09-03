"""Formal V2 business document rendering from clean facts and snapshots."""

from io import BytesIO
from pathlib import Path
from decimal import Decimal

from docx import Document
from flask import render_template

from v2_domain import format_batch_numbers, format_container_requirement


def format_decimal_compact(value):
    """Render a Decimal without insignificant zeroes or scientific notation."""
    if value is None:
        return ""
    value = Decimal(str(value))
    text = format(value, "f")
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text or "0"


def normalize_weight_input(value, display_unit):
    unit = (display_unit or "KGS").upper()
    if unit not in {"KGS", "MT"}:
        raise ValueError("Weight display unit must be KGS or MT.")
    if value in (None, ""):
        return None
    amount = Decimal(str(value))
    return amount * Decimal("1000") if unit == "MT" else amount


def format_weight(kg, display_unit):
    if kg is None:
        return ""
    unit = (display_unit or "KGS").upper()
    value = Decimal(str(kg)) / Decimal("1000") if unit == "MT" else Decimal(str(kg))
    return f"{format_decimal_compact(value)}{unit}"


def format_value_unit(value, unit):
    if value is None:
        return ""
    return f"{format_decimal_compact(value)}{unit or ''}"


def document_context(pi, kind):
    net_weight = sum((item.quantity for item in pi.items), 0) * 1000
    return {
        "pi": pi, "kind": kind, "net_weight_kg": net_weight,
        "container_display": format_container_requirement(pi.container_type, pi.container_count),
        "format_batches": lambda item: format_batch_numbers([b.batch_number for b in item.batches]),
    }


def render_invoice_html(pi, kind):
    return render_template("v2/documents/invoice.html", **document_context(pi, kind))


_LINE_BREAK = "\uf000"


def _replacement_text(value):
    return str(value or "").replace(";", _LINE_BREAK).replace("；", _LINE_BREAK)


def _replace_paragraph(paragraph, mapping):
    text = "".join(run.text for run in paragraph.runs)
    updated = text
    for key, value in mapping.items():
        updated = updated.replace(key, _replacement_text(value))
    if updated != text:
        if paragraph.runs:
            first = paragraph.runs[0]
            parts = updated.split(_LINE_BREAK)
            first.text = parts[0]
            for part in parts[1:]:
                first.add_break()
                first.add_text(part)
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = updated


def _replace_booking_table(table, mapping):
    """Replace booking cells, leaving the table-only Chinese reference blank."""
    for row in table.rows:
        row_mapping = mapping
        if any(cell.text.strip() == "约号" for cell in row.cells):
            row_mapping = dict(mapping)
            row_mapping["{{contract_number}}"] = ""
            row_mapping["{{ contract_number }}"] = ""
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_paragraph(paragraph, row_mapping)


def render_booking_docx(pi, template_path):
    quantity = sum((item.quantity for item in pi.items), 0)
    product_entries = []
    for item in pi.items:
        category = item.product_category_snapshot or ""
        brand_model = " ".join(filter(None, (item.product_brand_snapshot, item.product_model_snapshot)))
        product_entries.append(_LINE_BREAK.join(part for part in (category, brand_model) if part))
    products = _LINE_BREAK.join(entry for entry in product_entries if entry)
    notify_name = pi.customer_name_snapshot if pi.notify_party_same_as_consignee is not False else pi.notify_party_name_snapshot
    notify_address = pi.customer_address_snapshot if pi.notify_party_same_as_consignee is not False else pi.notify_party_address_snapshot
    notify_tax = pi.customer_tax_code_snapshot if pi.notify_party_same_as_consignee is not False else pi.notify_party_tax_code_snapshot
    product_hs = " / ".join(filter(None, [
        f"{item.product_model_snapshot or 'Product'}: {item.product_hs_code_snapshot}"
        for item in pi.items if item.product_hs_code_snapshot
    ]))
    mapping = {
        "{{ pi_no }}": pi.pi_no, "{{ quantity }}": format_decimal_compact(quantity),
        "{{vessel_info}}": pi.vessel_info, "{{ vessel_info }}": pi.vessel_info, "{ {vessel_info}}": pi.vessel_info,
        "{{booking_number}}": pi.booking_number, "{{ booking_number }}": pi.booking_number,
        "{{pi_no}}": pi.pi_no, "{{quantity}}": format_decimal_compact(quantity),
        "{{exporter_name}}": pi.exporter_name_snapshot, "{{exporter_address}}": pi.exporter_address_snapshot,
        "{{exporter_tax_code}}": pi.exporter_tax_code_snapshot,
        "{{customer_name}}": pi.customer_name_snapshot, "{{customer_address}}": pi.customer_address_snapshot,
        "{{customer_tax_code}}": pi.customer_tax_code_snapshot,
        "{{loading_port}}": pi.loading_port, "{{destination_port}}": pi.destination_port,
        "{{container_type_quantity}}": format_container_requirement(pi.container_type, pi.container_count),
        "{{shipping_mark}}": pi.shipping_mark, "{{product_category}}": products,
        "{{product_brand}}": "", "{{product_model}}": "",
        "{{freight_term}}": pi.freight_term, "{{contract_number}}": pi.contract_number or pi.pi_no,
        "{{freight_clause}}": pi.freight_clause, "{{waybill_option}}": pi.waybill_option,
        "{{exporter_ name }}": pi.exporter_name_snapshot,
        "{{exporter_ address }}": pi.exporter_address_snapshot,
        "{{exporter_ tax_code }}": pi.exporter_tax_code_snapshot,
        "{{customer _ name }}": pi.customer_name_snapshot,
        "{{ customer _address }}": pi.customer_address_snapshot,
        "{{ customer_tax_code }}": pi.customer_tax_code_snapshot,
        "{{notify_party_name}}": notify_name, "{{notify_party_address}}": notify_address,
        "{{notify_party_tax_code}}": notify_tax,
        "{{ loading_port }}": pi.loading_port, "{{ destination_port }}": pi.destination_port,
        "{{ container_type_quantity }}": format_container_requirement(pi.container_type, pi.container_count),
        "{{ shipping_mark }}": pi.shipping_mark, "{{ product_ category }} {{ product_brand }} {{ product_model }}": products,
        "{{ freight_term }}": pi.freight_term, "{{ contract_number }}": pi.contract_number or pi.pi_no,
        "{{ product_hs_codes }}": product_hs, "{{product_hs_codes}}": product_hs,
        "{ { freight_clause }}": pi.freight_clause, "{{ waybill_option }}": pi.waybill_option,
        "{{ container_number }}": pi.container_number, "{{ seal_number }}": pi.seal_number,
        "{{container_number}}": pi.container_number, "{{seal_number}}": pi.seal_number,
        "{{ container_type }}": pi.container_type, "{{container_type}}": pi.container_type,
        "{{ quantity_units }}": format_value_unit(pi.package_count, pi.package_unit),
        "{{quantity_units}}": format_value_unit(pi.package_count, pi.package_unit),
        "{{ gross_weight }}": format_weight(pi.gross_weight_kg, pi.gross_weight_display_unit),
        "{{gross_weight}}": format_weight(pi.gross_weight_kg, pi.gross_weight_display_unit),
        "{{ volume }}": format_value_unit(pi.volume_cbm, "CBM"), "{{volume}}": format_value_unit(pi.volume_cbm, "CBM"),
        "{{ vgm }}": format_weight(pi.vgm_kg, pi.vgm_display_unit), "{{vgm}}": format_weight(pi.vgm_kg, pi.vgm_display_unit),
        "{{ total_quantity }}": format_value_unit(pi.package_count, pi.package_unit), "{{ total_quantity_unit }}": "",
        "{{total_quantity}}": format_value_unit(pi.package_count, pi.package_unit), "{{total_quantity_unit}}": "",
        "{{ total_weight }}": format_weight(pi.gross_weight_kg, pi.gross_weight_display_unit), "{{ total_weight_unit }}": "",
        "{{total_weight}}": format_weight(pi.gross_weight_kg, pi.gross_weight_display_unit), "{{total_weight_unit}}": "",
        "{{ total_volume }}": format_value_unit(pi.volume_cbm, "CBM"), "{{ total_volume_unit }}": "",
        "{{total_volume}}": format_value_unit(pi.volume_cbm, "CBM"), "{{total_volume_unit}}": "",
        "{{ total_vgm }}": format_weight(pi.vgm_kg, pi.vgm_display_unit), "{{total_vgm}}": format_weight(pi.vgm_kg, pi.vgm_display_unit),
    }
    doc = Document(str(template_path))
    for paragraph in doc.paragraphs:
        _replace_paragraph(paragraph, mapping)
    for table in doc.tables:
        _replace_booking_table(table, mapping)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def booking_missing_fields(pi):
    """Hard requirements for an initial booking; post-loading facts are excluded."""
    fields = {
        "Vessel / Voyage": pi.vessel_info,
        "Container Type": pi.container_type,
        "Container Count": pi.container_count,
        "Shipping Mark": pi.shipping_mark,
        "Freight Term": pi.freight_term,
        "Waybill Option": pi.waybill_option,
    }
    return [label for label, value in fields.items() if value in (None, "")]
