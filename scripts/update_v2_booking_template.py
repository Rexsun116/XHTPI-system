#!/usr/bin/env python3
"""Build the V2-only Booking template without modifying V1's template."""
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
V1_PATH = ROOT / "templates" / "word" / "BN-Sample.docx"
V2_PATH = ROOT / "v2" / "templates" / "word" / "BN-Sample.docx"

def replace(paragraph, old, new):
    text = "".join(run.text for run in paragraph.runs)
    if old not in text:
        return
    text = text.replace(old, new)
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""

V2_PATH.parent.mkdir(parents=True, exist_ok=True)
doc = Document(V1_PATH)
for table in doc.tables:
    for row in table.rows:
        if row.cells and row.cells[0].text.strip() == "Notify Party":
            for cell in row.cells[1:]:
                for paragraph in cell.paragraphs:
                    replace(paragraph, "{{customer_name}}", "{{notify_party_name}}")
                    replace(paragraph, "{{customer_address}}", "{{notify_party_address}}")
                    replace(paragraph, "{{customer_tax_code}}", "{{notify_party_tax_code}}")
        if row.cells and "Description of Goods" in row.cells[0].text:
            for cell in row.cells[1:]:
                for paragraph in cell.paragraphs:
                    replace(paragraph, "HS CODE: 320611", "HS CODE: {{product_hs_codes}}")
doc.save(V2_PATH)
