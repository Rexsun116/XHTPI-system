"""Fresh V2 lineage, boot, clean facts, correction, and synthetic E2E tests."""

from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
import os
from pathlib import Path
import sqlite3
import subprocess
import sys
import tempfile
from unittest import TestCase

from werkzeug.security import generate_password_hash
from docx import Document

from order_lifecycle import LifecyclePolicyError, OrderModule, lifecycle_context
from v2.app import create_app
from v2.models import (
    BankAccount, Customer, Exporter, Factory, FreightForwarder, FreightQuote,
    FreightSettlement, OrderCorrectionSession, OrderFreightAgreement, OrderTask,
    PI, PIItem, Product, ProductBatch, TaskActivity, User, db,
)
from v2.services import (
    assert_correction_allows, close_correction_session, create_freight_agreement,
    freight_agreement_difference, open_correction_session,
    reconcile_order_tasks_for_pi, save_order_with_reconcile,
)


class V2CleanBaselineTest(TestCase):
    @classmethod
    def setUpClass(cls):
        cls.temp = tempfile.TemporaryDirectory(prefix="xhtpi-v2-tests-")
        cls.path = Path(cls.temp.name) / "v2.db"
        root = Path(__file__).resolve().parents[1]
        subprocess.run(
            [sys.executable, str(root / "scripts/init_v2_test_db.py"), str(cls.path)],
            cwd=root, check=True, capture_output=True, text=True,
        )
        cls.app = create_app(f"sqlite:///{cls.path}", testing=True)

    @classmethod
    def tearDownClass(cls):
        cls.temp.cleanup()

    def setUp(self):
        self.ctx = self.app.app_context(); self.ctx.push()
        for model in (TaskActivity, OrderTask, OrderCorrectionSession, ProductBatch,
                      OrderFreightAgreement, FreightSettlement, PIItem, PI, FreightQuote,
                      BankAccount, Product, FreightForwarder, Factory, Exporter, Customer, User):
            db.session.query(model).delete()
        db.session.commit()
        self.user = User(username="v2", password_hash=generate_password_hash("password"))
        self.customer = Customer(code="C1", name="Rahul", country="IN")
        self.exporter = Exporter(code="E1", name="AEA", country="HK")
        self.factory = Factory(code="F1", name="XHT", country="CN")
        self.forwarder = FreightForwarder(code="FF1", name="SeaLink", country="CN")
        self.product = Product(code="P1", model="R504", packaging="25kg bags")
        self.bank = BankAccount(
            code="BANK1", name="USD Account", beneficiary_name="AEA Limited",
            bank_name="Example Bank", account_number="123456", swift_code="EXAMPLEHK",
            currency="USD",
        )
        db.session.add_all([self.user, self.customer, self.exporter, self.factory,
                            self.forwarder, self.product, self.bank]); db.session.commit()

    def tearDown(self):
        db.session.rollback(); db.session.remove(); self.ctx.pop()

    def login_client(self):
        client = self.app.test_client()
        self.assertEqual(client.post("/login", data={"username": "v2", "password": "password"}).status_code, 302)
        return client

    def sales_payload(self, pi_no="V2-SALES"):
        return {
            "pi_no": pi_no, "pi_date": "2026-08-21", "order_type": "SALES",
            "customer_id": self.customer.id, "exporter_id": self.exporter.id,
            "bank_account_id": self.bank.id, "currency": "USD",
            "payment_terms": "20% advance", "advance_payment_percent": "20", "planned_shipment_date": "2026-09-20",
            "items": [{"product_id": self.product.id, "unit_price": "10.00", "quantity": "100"}],
        }

    def test_01_fresh_migration_fk_integrity_and_no_legacy_columns(self):
        with sqlite3.connect(self.path) as conn:
            conn.execute("PRAGMA foreign_keys=ON")
            self.assertEqual(conn.execute("PRAGMA integrity_check").fetchone()[0], "ok")
            self.assertEqual(conn.execute("PRAGMA foreign_key_check").fetchall(), [])
            self.assertEqual(conn.execute("SELECT version_num FROM alembic_version").fetchone()[0], "v2_0003")
            columns = {row[1] for row in conn.execute("PRAGMA table_info(pi)")}
        self.assertNotIn("commission_exporter_id", columns)
        self.assertNotIn("payment_received", columns)
        self.assertNotIn("ocean_freight", columns)

    def test_02_app_boot_login_create_sales_and_targeted_reminder(self):
        client = self.login_client()
        response = client.post("/orders", json=self.sales_payload())
        self.assertEqual(response.status_code, 201)
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "V2-SALES"))
        self.assertEqual(pi.advance_payment_amount, Decimal("200.00"))
        self.assertEqual(pi.balance_payment_amount, Decimal("800.00"))
        task = db.session.scalar(db.select(OrderTask).where(OrderTask.pi_id == pi.id))
        self.assertEqual((task.task_code, task.status), ("PAYMENT_ADVANCE_WAITING", "WAITING"))
        self.assertEqual(pi.bank_account_number_snapshot, "123456")

    def test_03_bank_master_change_does_not_change_snapshot(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("BANK-SNAP"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "BANK-SNAP"))
        self.bank.account_number = "999999"; db.session.commit(); db.session.refresh(pi)
        self.assertEqual(pi.bank_account_number_snapshot, "123456")

    def test_04_product_batches_are_item_scoped_and_cascade(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("BATCH"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "BATCH")); item = pi.items[0]
        item.batches.extend([ProductBatch(batch_number="A", display_order=1), ProductBatch(batch_number="B", display_order=2)])
        db.session.commit(); item_id = item.id
        self.assertEqual([b.batch_number for b in item.batches], ["A", "B"])
        db.session.delete(item); db.session.commit()
        self.assertEqual(db.session.scalar(db.select(db.func.count()).select_from(ProductBatch).where(ProductBatch.pi_item_id == item_id)), 0)

    def test_05_agreement_snapshot_and_usd_difference_without_cny_comparison(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("FREIGHT"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "FREIGHT"))
        quote = FreightQuote(
            freight_forwarder_id=self.forwarder.id, departure_port="Shanghai",
            destination_port="Mumbai", amount=Decimal("100"), currency="USD",
        )
        db.session.add(quote); db.session.flush()
        agreement = create_freight_agreement(pi, quote); db.session.add(agreement)
        settlement = FreightSettlement(
            pi_id=pi.id, usd_bill_required=True, usd_bill_amount=Decimal("120"),
            cny_bill_required=True, cny_bill_amount=Decimal("800"),
        )
        db.session.add(settlement); db.session.flush()
        result = freight_agreement_difference(agreement, settlement)
        self.assertEqual(result["difference"], Decimal("20.00"))
        reconcile_order_tasks_for_pi(pi); db.session.commit()
        task = db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "FREIGHT_BILL_DIFFERS_FROM_AGREED_QUOTE"))
        self.assertEqual(task.context_payload["actual_amount"], "120.00")
        self.assertNotIn("800", str(task.context_payload))

    def test_06_commission_derived_and_override_requires_reason(self):
        payload = self.sales_payload("COMMISSION"); payload.update({
            "order_type": "COMMISSION", "exporter_id": None,
            "commission_factory_id": self.factory.id, "commission_rate": "2.5",
            "commission_currency": "USD",
        })
        response = self.login_client().post("/orders", json=payload)
        self.assertEqual(response.status_code, 201)
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "COMMISSION"))
        self.assertEqual(pi.commission_amount, Decimal("25.00"))
        self.assertFalse(hasattr(pi, "commission_exporter_id"))
        pi.commission_amount_mode = "EXPLICIT_OVERRIDE"; pi.commission_amount = Decimal("30")
        with self.assertRaises(ValueError):
            pi.derive_commission()
        pi.commission_override_reason = "Contractually agreed"; self.assertEqual(pi.derive_commission(), Decimal("30"))

    def test_07_completed_correction_is_module_scoped_and_audited(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("CORRECT"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "CORRECT")); pi.status = "COMPLETED"; db.session.commit()
        with self.assertRaises(LifecyclePolicyError):
            assert_correction_allows(pi, "PAYMENT_PLAN")
        with self.assertRaises(LifecyclePolicyError):
            open_correction_session(pi, "PAYMENT", "", self.user.id)
        correction = open_correction_session(pi, "PAYMENT", "Correct receipt", self.user.id)
        self.assertEqual(correction.opened_by_id, self.user.id)
        assert_correction_allows(pi, "PAYMENT_PLAN")
        with self.assertRaises(LifecyclePolicyError):
            assert_correction_allows(pi, "FREIGHT_SETTLEMENT")
        close_correction_session(correction, self.user.id, note="Corrected")
        self.assertIsNotNone(correction.closed_at)

    def test_08_dashboard_get_is_read_only_and_clean_has_no_legacy_import(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("DASH"))
        before = (db.session.scalar(db.select(db.func.count()).select_from(OrderTask)),
                  db.session.scalar(db.select(db.func.count()).select_from(TaskActivity)))
        self.assertEqual(client.get("/v2/").status_code, 200)
        after = (db.session.scalar(db.select(db.func.count()).select_from(OrderTask)),
                 db.session.scalar(db.select(db.func.count()).select_from(TaskActivity)))
        self.assertEqual(before, after)
        self.assertFalse(any(task.task_code.startswith("LEGACY") for task in db.session.scalars(db.select(OrderTask))))

    def test_09_lifecycle_matrix_completed_and_shipped(self):
        self.assertTrue(lifecycle_context("SHIPPED").can_edit(OrderModule.FREIGHT_SETTLEMENT))
        self.assertFalse(lifecycle_context("COMPLETED").can_edit(OrderModule.FREIGHT_SETTLEMENT))
        self.assertTrue(lifecycle_context("COMPLETED", correction_modules={OrderModule.FREIGHT_SETTLEMENT}).can_edit(OrderModule.FREIGHT_SETTLEMENT))

    def test_10_synthetic_sales_lifecycle_and_reminder_resolution(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("E2E"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "E2E"))
        pi.status = "PRE_SHIPMENT"
        pi.container_type, pi.container_count = "20GP", 1
        pi.container_loading_at = datetime(2026, 8, 21, 9, 0)
        pi.coo_required = True
        pi.etd, pi.eta = date(2026, 8, 21), date(2026, 8, 22)
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 21, 10, 0)); db.session.commit()
        driver = db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "SHIPPING_DRIVER_INFO"))
        departure = db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "SHIPPING_ACTUAL_DEPARTURE"))
        self.assertEqual((driver.health, departure.health), ("EXCEPTION", "EXCEPTION"))
        self.assertIsNotNone(db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "DOCUMENT_COO")))

        pi.driver_name, pi.driver_phone, pi.vehicle_number = "Li", "13800000000", "沪A12345"
        pi.actual_departure_date = date(2026, 8, 21); pi.status = "SHIPPED"
        pi.advance_received_amount = pi.advance_payment_amount
        settlement = FreightSettlement(pi_id=pi.id, usd_bill_required=True, cny_bill_required=True)
        db.session.add(settlement); db.session.flush()
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 21, 12, 0)); db.session.commit()
        self.assertEqual(driver.status, "DONE"); self.assertEqual(departure.status, "DONE")
        advance = db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "PAYMENT_ADVANCE_WAITING"))
        self.assertEqual(advance.status, "DONE")
        self.assertIsNotNone(db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "FREIGHT_USD_AMOUNT_CAPTURE")))
        self.assertIsNotNone(db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "FREIGHT_CNY_AMOUNT_CAPTURE")))

        settlement.usd_bill_amount, settlement.cny_bill_amount = Decimal("120"), Decimal("800")
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 21, 13, 0)); db.session.commit()
        self.assertIsNotNone(db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "FREIGHT_USD_AMOUNT_CONFIRM")))
        settlement.usd_bill_confirmed = settlement.cny_bill_confirmed = True
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 21, 14, 0)); db.session.commit()
        self.assertIsNotNone(db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "FREIGHT_INVOICE_ISSUED")))
        settlement.invoice_issued = True
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 21, 15, 0)); db.session.commit()
        payment_task = db.session.scalar(db.select(OrderTask).where(OrderTask.task_code == "FREIGHT_PAYMENT_CONFIRM"))
        self.assertEqual(payment_task.status, "ACTION")
        settlement.payment_status = "PAID"
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 21, 16, 0)); db.session.commit()
        self.assertEqual(payment_task.status, "DONE")

        pi.status = "ARRIVED"; pi.actual_arrival_date = date(2026, 8, 22)
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 8, 22, 18, 0)); db.session.commit()
        pi.status = "COMPLETED"; db.session.commit()
        self.assertFalse(lifecycle_context(pi).can_edit(OrderModule.COMMERCIAL_CORE))

    def test_11_user_facing_master_and_multi_item_order_ui(self):
        client=self.login_client()
        response=client.post("/v2/master/customers",data={"code":"C2","name":"Second Customer"})
        self.assertEqual(response.status_code,302)
        form={"pi_no":"UI-ORDER","pi_date":"2026-08-21","order_type":"SALES",
            "customer_id":str(self.customer.id),"exporter_id":str(self.exporter.id),
            "bank_account_id":str(self.bank.id),"currency":"USD","advance_payment_percent":"20",
            "payment_terms":"20% advance","planned_shipment_date":"2026-09-20",
            "product_0":str(self.product.id),"factory_0":str(self.factory.id),"unit_price_0":"10","quantity_0":"10","quantity_unit_0":"MT",
            "product_1":str(self.product.id),"factory_1":str(self.factory.id),"unit_price_1":"20","quantity_1":"5","quantity_unit_1":"MT"}
        response=client.post("/v2/orders/new",data=form)
        self.assertEqual(response.status_code,302)
        pi=db.session.scalar(db.select(PI).where(PI.pi_no=="UI-ORDER"))
        self.assertEqual((len(pi.items),pi.contract_total),(2,Decimal("200.00")))
        self.assertIn("Order Control Center",client.get("/v2/").get_data(as_text=True))

    def test_12_documents_generate_in_v2_directory(self):
        client=self.login_client(); client.post("/orders",json=self.sales_payload("DOCS"))
        pi=db.session.scalar(db.select(PI).where(PI.pi_no=="DOCS"))
        pi.vessel_info="TEST VESSEL"; pi.container_type="20GP"; pi.container_count=1
        pi.shipping_mark="TEST MARK"; pi.freight_term="FOB"; pi.waybill_option="ORIGINAL"; db.session.commit()
        for kind,magic in (("pi",b"%PDF"),("invoice",b"%PDF"),("packing",b"%PDF"),("booking",b"PK")):
            response=client.get(f"/v2/orders/{pi.id}/documents/{kind}")
            self.assertEqual(response.status_code,200); self.assertTrue(response.data.startswith(magic))

    def test_13_v2_path_guard_and_csrf(self):
        root=Path(__file__).resolve().parents[1]
        with self.assertRaises(RuntimeError): create_app(f"sqlite:///{root/'instance'/'database.db'}",testing=True)
        secure=create_app(f"sqlite:///{self.path}",testing=False,secret_key="a-secure-test-secret")
        self.assertEqual(secure.test_client().post("/login",data={"username":"v2","password":"password"}).status_code,400)

    def test_14_product_hs_snapshot_and_booking_never_uses_fixed_code(self):
        self.product.hs_code = "390761"; db.session.commit()
        client = self.login_client(); client.post("/orders", json=self.sales_payload("HS-SNAP"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "HS-SNAP"))
        self.assertEqual(pi.items[0].product_hs_code_snapshot, "390761")
        self.product.hs_code = "999999"; db.session.commit()
        pi.vessel_info = "Vessel 1"; pi.container_type = "20GP"; pi.container_count = 1
        pi.shipping_mark = "MARK"; pi.freight_term = "FOB"; pi.waybill_option = "ORIGINAL"
        db.session.commit()
        response = client.get(f"/v2/orders/{pi.id}/documents/booking")
        self.assertEqual(response.status_code, 200)
        doc = Document(BytesIO(response.data))
        text = "\n".join(p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs)
        self.assertIn("390761", text)
        self.assertNotIn("320611", text)

    def test_15_booking_uses_consignee_or_custom_notify_snapshot(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("NOTIFY"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "NOTIFY"))
        pi.customer_name_snapshot = "Consignee"; pi.customer_address_snapshot = "Consignee Address"
        pi.vessel_info = "Vessel 1"; pi.container_type = "20GP"; pi.container_count = 1
        pi.shipping_mark = "MARK"; pi.freight_term = "FOB"; pi.waybill_option = "ORIGINAL"
        db.session.commit()
        response = client.get(f"/v2/orders/{pi.id}/documents/booking")
        doc = Document(BytesIO(response.data)); text = "\n".join(p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs)
        self.assertIn("Consignee", text)
        pi.notify_party_same_as_consignee = False; pi.notify_party_name_snapshot = "Custom Notify"
        pi.notify_party_address_snapshot = "Custom Address"; db.session.commit()
        response = client.get(f"/v2/orders/{pi.id}/documents/booking")
        doc = Document(BytesIO(response.data)); text = "\n".join(p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs)
        self.assertIn("Custom Notify", text)

    def test_16_v1_booking_template_is_isolated_from_v2_template(self):
        root = Path(__file__).resolve().parents[1]
        def text_for(path):
            doc = Document(path)
            return "\n".join(p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs)
        self.assertIn("HS CODE: 320611", text_for(root / "templates" / "word" / "BN-Sample.docx"))
        self.assertIn("{{product_hs_codes}}", text_for(root / "v2" / "templates" / "word" / "BN-Sample.docx"))

    def test_17_pre_shipment_document_and_item_hs_editors_are_controlled(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("PRE-EDIT"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "PRE-EDIT")); pi.status = "PRE_SHIPMENT"
        db.session.commit()
        response = client.post(f"/v2/orders/{pi.id}/document-requirements", data={"coc_required": "true"})
        self.assertEqual(response.status_code, 302); self.assertTrue(pi.coc_required)
        response = client.post(f"/v2/orders/{pi.id}/booking-hs-codes", data={f"product_hs_code_snapshot_{pi.items[0].id}": "281700"})
        self.assertEqual(response.status_code, 302)
        self.assertEqual(pi.items[0].product_hs_code_snapshot, "281700")
        self.assertIsNone(self.product.hs_code)

    def test_18_legacy_loading_datetime_is_read_fallback_only(self):
        client = self.login_client(); client.post("/orders", json=self.sales_payload("LOAD-FALLBACK"))
        pi = db.session.scalar(db.select(PI).where(PI.pi_no == "LOAD-FALLBACK"))
        pi.status = "PRE_SHIPMENT"; pi.container_loading_at = datetime(2026, 9, 20, 14, 0)
        pi.advance_received_amount = pi.advance_payment_amount
        reconcile_order_tasks_for_pi(pi, now=datetime(2026, 9, 20, 10)); db.session.commit()
        task = db.session.scalar(db.select(OrderTask).where(OrderTask.pi_id == pi.id, OrderTask.task_code == "SHIPPING_CONTAINER_LOADING"))
        self.assertIsNone(task)
        self.assertIsNone(pi.container_loading_date)
        self.assertIsNone(pi.container_loading_period)
