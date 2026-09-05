"""Behavioral parity tests for the V2 local UAT gate."""

from datetime import date, datetime, timedelta
from decimal import Decimal
from io import BytesIO
import tempfile
from pathlib import Path
from unittest import TestCase

from docx import Document

from werkzeug.security import generate_password_hash

from v2.app import create_app
from v2.models import (Customer, Exporter, FreightSettlement, OrderFreightAgreement,
                       OrderTask, PI, PIItem, ProductBatch, TaskActivity, User, db)
from v2.documents import format_decimal_compact, render_booking_docx, render_invoice_html
from v2.selector import select_next_action
from v2.services import reconcile_order_tasks_for_pi
from v2.task_service import TaskOperationError, mark_done, move_to_waiting, reopen


class V2UATGateTest(TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        uri = f"sqlite:///{Path(self.tmp.name) / 'v2.db'}"
        self.app = create_app(uri, testing=True)
        self.ctx = self.app.app_context(); self.ctx.push(); db.create_all()
        self.user = User(username="uat", password_hash=generate_password_hash("secret"))
        self.customer = Customer(code="C", name="Customer")
        self.exporter = Exporter(code="E", name="Exporter")
        db.session.add_all([self.user, self.customer, self.exporter]); db.session.commit()

    def tearDown(self):
        db.session.remove(); self.ctx.pop(); self.tmp.cleanup()

    def pi(self, **values):
        defaults = dict(pi_no=f"UAT-{PI.query.count()+1}", order_type="SALES", status="NEW",
            pi_date=date(2026,8,21), customer_id=self.customer.id, exporter_id=self.exporter.id,
            customer_name_snapshot=self.customer.name, exporter_name_snapshot=self.exporter.name,
            currency="USD", advance_payment_amount=Decimal("200"), balance_payment_amount=Decimal("800"))
        defaults.update(values); pi=PI(**defaults); pi.items.append(PIItem(unit_price=Decimal("100"),quantity=Decimal("10"),
            quantity_unit="MT",line_total=Decimal("1000"),product_model_snapshot="R504")); db.session.add(pi); db.session.flush(); return pi

    def task(self, pi, code):
        return db.session.scalar(db.select(OrderTask).where(OrderTask.pi_id==pi.id,OrderTask.task_code==code))

    def client(self):
        client=self.app.test_client()
        with client.session_transaction() as session: session["_user_id"]=str(self.user.id); session["_fresh"]=True
        return client

    def test_document_rules_and_requirement_cancel_reactivate(self):
        now=datetime(2026,8,21,12); pi=self.pi(status="PRE_SHIPMENT",container_loading_at=now,
            coo_required=True,apta_required=True,export_license_required=True,customs_docs_required=True,
            coc_required=True,coa_required=True,original_bl_required=True,obd_electronic_required=True,
            insurance_original_required=True,insurance_electronic_required=True)
        reconcile_order_tasks_for_pi(pi,now=now)
        for code in ("DOCUMENT_COO","DOCUMENT_APTA","DOCUMENT_EXPORT_LICENSE","DOCUMENT_CUSTOMS","DOCUMENT_COC","DOCUMENT_COA"):
            self.assertEqual(self.task(pi,code).status,"ACTION")
        self.assertEqual(self.task(pi,"DOCUMENT_ORIGINAL_BL").status,"UPCOMING")
        pi.coo_required=False; reconcile_order_tasks_for_pi(pi,now=now)
        self.assertEqual(self.task(pi,"DOCUMENT_COO").status,"CANCELLED")
        task_id=self.task(pi,"DOCUMENT_COO").id
        pi.coo_required=True; reconcile_order_tasks_for_pi(pi,now=now)
        self.assertEqual(self.task(pi,"DOCUMENT_COO").id,task_id)
        self.assertEqual(self.task(pi,"DOCUMENT_COO").status,"ACTION")

    def test_original_mail_required_input_and_balance_earliest_plus_three(self):
        now=datetime(2026,8,21,10); pi=self.pi(status="SHIPPED",original_bl_required=True,
            insurance_original_required=False,original_documents_mail_required=True)
        reconcile_order_tasks_for_pi(pi,now=now)
        bl=self.task(pi,"DOCUMENT_ORIGINAL_BL"); mark_done(bl,self.user.id); reconcile_order_tasks_for_pi(pi,now=now)
        mail=self.task(pi,"ORIGINAL_DOCUMENTS_MAIL"); self.assertEqual(mail.status,"ACTION")
        with self.assertRaises(TaskOperationError): mark_done(mail,self.user.id,payload={})
        mark_done(mail,self.user.id,payload={"tracking_number":"DHL-1","carrier":"DHL"})
        email=self.task(pi,"PAYMENT_EMAIL"); email.completed_at=now-timedelta(days=1); email.status="DONE"
        mail.completed_at=now-timedelta(days=4)
        reconcile_order_tasks_for_pi(pi,now=now)
        follow=self.task(pi,"PAYMENT_BALANCE_FOLLOWUP")
        self.assertEqual(follow.status,"ACTION")
        self.assertEqual(follow.context_payload["trigger_date"],"2026-08-20")
        completed=[a for a in mail.activities if a.event_type=="COMPLETED"][-1]
        self.assertEqual(completed.payload["tracking_number"],"DHL-1")

    def test_waiting_due_reconcile_and_history_preserved(self):
        now=datetime(2026,8,21,10); pi=self.pi(status="SHIPPED")
        reconcile_order_tasks_for_pi(pi,now=now); email=self.task(pi,"PAYMENT_EMAIL")
        mark_done(email,self.user.id); email.completed_at=now-timedelta(days=4); reconcile_order_tasks_for_pi(pi,now=now)
        follow=self.task(pi,"PAYMENT_BALANCE_FOLLOWUP")
        move_to_waiting(follow,self.user.id,waiting_on="CUSTOMER",next_follow_up_at=now-timedelta(hours=1),note="sent")
        reconcile_order_tasks_for_pi(pi,now=now)
        self.assertEqual(follow.status,"ACTION"); self.assertEqual(follow.health,"OVERDUE")
        self.assertEqual(len([a for a in follow.activities if a.event_type=="WAITING_STARTED"]),1)

    def test_rule_data_reverse_reactivates_without_duplicates(self):
        now=datetime(2026,8,21,10); pi=self.pi(status="PRE_SHIPMENT",etd=date(2026,8,20))
        reconcile_order_tasks_for_pi(pi,now=now); task=self.task(pi,"SHIPPING_ACTUAL_DEPARTURE"); task_id=task.id
        pi.actual_departure_date=date(2026,8,21); reconcile_order_tasks_for_pi(pi,now=now); self.assertEqual(task.status,"DONE")
        pi.actual_departure_date=None; reconcile_order_tasks_for_pi(pi,now=now); self.assertEqual(task.status,"ACTION")
        count=len(task.activities); reconcile_order_tasks_for_pi(pi,now=now)
        self.assertEqual(task.id,task_id); self.assertEqual(len(task.activities),count)

    def test_freight_capture_confirm_snapshot_difference_and_payment(self):
        now=datetime(2026,8,21,10); pi=self.pi(status="SHIPPED",actual_departure_date=date(2026,8,10))
        settlement=FreightSettlement(pi_id=pi.id,usd_bill_required=True,cny_bill_required=True)
        agreement=OrderFreightAgreement(pi_id=pi.id,freight_forwarder_name_snapshot="FF",amount=Decimal("100"),currency="USD")
        db.session.add_all([settlement,agreement]); reconcile_order_tasks_for_pi(pi,now=now)
        self.assertEqual(self.task(pi,"FREIGHT_USD_AMOUNT_CAPTURE").status,"ACTION")
        self.assertEqual(self.task(pi,"FREIGHT_CNY_AMOUNT_CAPTURE").status,"ACTION")
        settlement.usd_bill_amount=Decimal("120"); settlement.cny_bill_amount=Decimal("700")
        reconcile_order_tasks_for_pi(pi,now=now)
        self.assertEqual(self.task(pi,"FREIGHT_BILL_DIFFERS_FROM_AGREED_QUOTE").health,"EXCEPTION")
        self.assertEqual(self.task(pi,"FREIGHT_USD_AMOUNT_CONFIRM").context_payload["amount"],"120.00")

    def test_next_action_priority_and_dashboard_is_read_only(self):
        pi=self.pi(); exception=OrderTask(pi_id=pi.id,task_code="E",title="E",source="AUTO",status="ACTION",health="EXCEPTION",completion_mode="RULE_DATA",priority=100)
        action=OrderTask(pi_id=pi.id,task_code="A",title="A",source="AUTO",status="ACTION",health="NORMAL",completion_mode="MANUAL",priority=1)
        db.session.add_all([action,exception]); db.session.commit()
        self.assertIs(select_next_action([action,exception]),exception)
        before=(OrderTask.query.count(),TaskActivity.query.count())
        client=self.client()
        self.assertEqual(client.get("/v2/").status_code,200)
        self.assertEqual((OrderTask.query.count(),TaskActivity.query.count()),before)

    def test_reopen_preserves_completed_activity(self):
        pi=self.pi(); task=OrderTask(pi_id=pi.id,task_code="M",title="Manual",source="MANUAL",status="ACTION",health="NORMAL",completion_mode="MANUAL")
        db.session.add(task); db.session.flush(); mark_done(task,self.user.id); completed=next(a.id for a in task.activities if a.event_type=="COMPLETED")
        reopen(task,self.user.id,reason="correct it")
        self.assertEqual(task.status,"ACTION"); self.assertIn(completed,[a.id for a in task.activities])

    def test_required_input_http_and_rule_data_done_rejection(self):
        pi=self.pi(); required=OrderTask(pi_id=pi.id,task_code="ORIGINAL_DOCUMENTS_MAIL",title="Mail",source="AUTO",
            status="ACTION",health="NORMAL",completion_mode="MANUAL_REQUIRED_INPUT")
        rule=OrderTask(pi_id=pi.id,task_code="RULE",title="Rule",source="AUTO",status="ACTION",health="NORMAL",completion_mode="RULE_DATA")
        db.session.add_all([required,rule]); db.session.commit(); client=self.client()
        self.assertEqual(client.post(f"/v2/tasks/{required.id}/done",data={}).status_code,400)
        self.assertEqual(client.post(f"/v2/tasks/{required.id}/done",data={"tracking_number":"T-1","carrier":"DHL"}).status_code,302)
        self.assertEqual(client.post(f"/v2/tasks/{rule.id}/done",data={}).status_code,400)
        self.assertEqual(required.activities[-1].payload["tracking_number"],"T-1")

    def test_document_content_batches_and_booking_placeholders(self):
        pi=self.pi(container_type="20GP",container_count=1,loading_port="SHANGHAI",destination_port="MUMBAI",
                   package_count=10,package_unit="BAGS",gross_weight_kg=Decimal("10100"),volume_cbm=Decimal("20"))
        pi.items[0].batches.extend([ProductBatch(batch_number="A",display_order=0),ProductBatch(batch_number="B",display_order=1)])
        html=render_invoice_html(pi,"packing")
        self.assertIn("BATCH NO.",html); self.assertIn("A / B",html); self.assertIn("10100",html)
        root=Path(__file__).resolve().parents[1]
        data=render_booking_docx(pi,root/"v2"/"templates"/"word"/"BN-Sample.docx").getvalue()
        self.assertGreater(len(data),1000)

    def test_booking_compact_units_and_multiline_snapshots(self):
        pi=self.pi(container_type="20GP", container_count=1, vessel_info="V", shipping_mark="M", freight_term="FOB", waybill_option="ORIGINAL",
                   package_count=800, package_unit="BAGS", gross_weight_kg=Decimal("20000.000"), gross_weight_display_unit="MT",
                   volume_cbm=Decimal("25.500"), exporter_name_snapshot="ABC COMPANY;ROOM 1208；CHENGDU, CHINA",
                   customer_address_snapshot="LINE 1;LINE 2")
        root=Path(__file__).resolve().parents[1]
        doc=Document(render_booking_docx(pi,root/"v2"/"templates"/"word"/"BN-Sample.docx"))
        text="\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(p.text for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs)
        self.assertIn("800BAGS", text); self.assertIn("20MT", text); self.assertIn("25.5CBM", text)
        self.assertIn("ABC COMPANY\nROOM 1208\nCHENGDU, CHINA", text)
        self.assertEqual((format_decimal_compact(Decimal("20.000")), format_decimal_compact(Decimal("20.500")), format_decimal_compact(Decimal("20000.000"))), ("20", "20.5", "20000"))

    def test_booking_goods_are_item_snapshot_lines_and_table_reference_is_blank(self):
        pi=self.pi(pi_no="PI-REFERENCE", container_type="20GP", container_count=1, vessel_info="V", shipping_mark="M", freight_term="FOB", waybill_option="ORIGINAL",
                   package_count=960, package_unit="BAGS")
        pi.items[0].product_category_snapshot="TITANIUM DIOXIDE"; pi.items[0].product_brand_snapshot="BLR"; pi.items[0].product_model_snapshot="R504"
        pi.items.append(PIItem(unit_price=Decimal("1"), quantity=Decimal("2"), quantity_unit="MT", line_total=Decimal("2"),
                               product_category_snapshot="CARBON BLACK", product_brand_snapshot="CB", product_model_snapshot="N330"))
        root=Path(__file__).resolve().parents[1]
        doc=Document(render_booking_docx(pi,root/"v2"/"templates"/"word"/"BN-Sample.docx"))
        goods=doc.tables[0].rows[8].cells[1].paragraphs[0].text
        self.assertIn("TITANIUM DIOXIDE\nBLR R504\nCARBON BLACK\nCB N330", goods)
        self.assertIn("960BAGS", doc.tables[0].rows[13].cells[5].text)
        self.assertEqual(doc.tables[0].rows[9].cells[7].text.strip(), "")

    def test_booking_browser_route_returns_current_booking_content(self):
        pi=self.pi(status="PRE_SHIPMENT", pi_no="ROUTE-BOOKING", container_type="20GP", container_count=1,
                   vessel_info="V", shipping_mark="M", freight_term="FOB", waybill_option="ORIGINAL",
                   package_count=960, package_unit="BAGS", gross_weight_kg=Decimal("20000"),
                   gross_weight_display_unit="MT", volume_cbm=Decimal("25.5"))
        pi.items[0].product_category_snapshot="TITANIUM DIOXIDE"; pi.items[0].product_brand_snapshot="BLR"; pi.items[0].product_model_snapshot="R504"
        response=self.client().get(f"/v2/orders/{pi.id}/documents/booking")
        self.assertEqual(response.status_code, 200)
        doc=Document(BytesIO(response.data)); table=doc.tables[0]
        self.assertIn("960BAGS", table.rows[13].cells[5].text)
        self.assertEqual(table.rows[9].cells[7].text.strip(), "")
        self.assertIn("TITANIUM DIOXIDE\nBLR R504", table.rows[8].cells[1].text)

    def test_commission_override_reason_required(self):
        pi=self.pi(order_type="COMMISSION",commission_rate=Decimal("2.5"),commission_currency="USD",
                   commission_amount_mode="EXPLICIT_OVERRIDE",commission_amount=Decimal("30"))
        with self.assertRaises(ValueError): pi.derive_commission()
        pi.commission_override_reason="commercial exception"
        self.assertEqual(pi.derive_commission(),Decimal("30"))
